import os

from docx import Document
from docx.shared import Inches

from .gen_trace_matrix_base import GenTraceMatrixBase
from ..common.gen_base_docx import GenBaseDocx
from ... import services
from ...constants import Constants


# -------------------
## Generates a Trace Matrix report in docx format
class GenTraceMatrixDocx(GenBaseDocx, GenTraceMatrixBase):
    # -------------------
    ## constructor
    #
    # @param matrix   the data to use
    def __init__(self, matrix):
        GenBaseDocx.__init__(self)
        GenTraceMatrixBase.__init__(self, matrix)

        ## holds reference to the current table being generated
        self._tbl = None
        ## holds current reqmt id for the row
        self._row_id = None
        ## holds current reqmt info for the row
        self._row_info = None

    # -------------------
    ## generate the report
    #
    # @return None
    def gen(self):
        self._gen_init()

        ## see GenBaseDocx
        self._doc = Document()
        self._doc_init()

        self._gen_test_run_details()
        self._gen_title('Trace Matrix')

        self._gen_trace_table()

        self._build()

    # -------------------
    ## generate a trace matrix table
    #
    # @return None
    def _gen_trace_table(self):
        self._create_table()
        self._gen_table_header_row()

        self._load_matrix_requirements()
        ## see base class doc
        self._report_requirements(self._report_header, self._report_desc, self._report_details)

    # -------------------
    ## callback to report the current requirement and tester info
    #
    # @param req_id   the requirement id to report
    # @param info     the info to report
    # @return None
    def _report_header(self, req_id, info):
        if self._row_id is not None:
            row_cells = self._tbl.add_row().cells
            self._set_cell_id(row_cells[0])
            self._set_cell_info(row_cells[1])

        self._row_id = f'{req_id}\n{info}'
        self._row_info = []

    # -------------------
    ## callback to report the current requirement's description (if available)
    #
    # @param desc     the description to report
    # @return None
    def _report_desc(self, desc):
        self._row_info.append(desc)

    # -------------------
    ## callback to report protocol id and info for the current requirement
    #
    # @param proto_id     the protocol id to report
    # @param proto_info   the protocol info to report
    # @return None
    def _report_details(self, proto_id, proto_info):
        if proto_id is None and proto_info is None:
            row_cells = self._tbl.add_row().cells
            self._set_cell_id(row_cells[0])
            self._set_cell_info(row_cells[1])

            self._row_id = None
            self._row_info = None
        elif proto_id is None:
            self._row_info.append(proto_info)
        else:
            self._row_info.append(f'{proto_id} {proto_info}')

    # -------------------
    ## create a table with 2 columns
    #
    # @return None
    def _create_table(self):
        # TODO add 0.25 space before and after table
        #     t = Table(self._tbl,
        #               colWidths=tbl_widths,
        #               spaceBefore=0.25 * inch,
        #               spaceAfter=0.25 * inch,
        #               )
        self._tbl = self._doc.add_table(rows=0, cols=2)
        self._tbl.style = 'Table Grid'
        self._align_table_center(self._tbl)

        self._tbl.columns[0].width = Inches(0.75)  # reqid
        self._tbl.columns[1].width = Inches(6.25)  # proto_info

    # -------------------
    ## generate header2 row
    #
    # @return None
    def _gen_table_header_row(self):
        row = self._tbl.add_row().cells
        self._set_header_cell(row[0], 'Req.')
        self._set_header_cell(row[1], 'Protocol')

    # -------------------
    ## set header1 cell to the given text
    #
    # @param cell  the cell to set
    # @param text  the text to set the cell to
    # @return None
    def _set_header_cell(self, cell, text):
        cell.paragraphs[0].add_run(text)
        cell.paragraphs[0].style = 'ver_table1_header'
        self._set_cell_blue_background(cell)
        self._align_vertical_top(cell)

    # -------------------
    ## set the requirement id in the given cell
    #
    # @param cell  the cell to write to
    # @return None
    def _set_cell_id(self, cell):
        cell.text = ''
        cell.paragraphs[0].add_run(self._row_id)
        cell.paragraphs[0].style = 'ver_table1_cell'
        self._align_vertical_center(cell)
        cell.alignment = self._align_center()

    # -------------------
    ## set the protocol info in the given cell
    #
    # @param cell  the cell to write to
    # @return None
    def _set_cell_info(self, cell):
        cell.text = ''
        self._align_vertical_center(cell)
        cell.alignment = self._align_left()
        # self._row_info
        cell.paragraphs[0].style = 'ver_table1_cell'
        run = cell.paragraphs[0].add_run('Desc: ')
        run.italic = True
        run.bold = True
        run = cell.paragraphs[0].add_run(self._row_info[0])
        run.italic = True
        run.bold = False

        for i in range(1, len(self._row_info)):
            txt = self._row_info[i]
            if i == 1:
                txt = f'\n{txt}'

            boldness = False
            if txt.find('<b>') != -1 or txt.find('</b>') != -1 or txt.find('<b/>') != -1:
                boldness = True
                txt = txt.replace('<b>', '')
                txt = txt.replace('</b>', '')
                txt = txt.replace('<b/>', '')

            para = cell.add_paragraph(style='ver_table1_cell')
            run = para.add_run(txt)
            run.bold = boldness

    # -------------------
    ## build the document
    #
    # @return None
    def _build(self):
        path = os.path.join(services.cfg.outdir, f'{Constants.trace_fname}.docx')
        self._doc.save(path)
