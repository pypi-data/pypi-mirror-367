import docx
import docx.enum
import docx.enum.section
import docx.enum.style
import docx.enum.table
import docx.enum.text
import docx.shared
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

from ... import services
from ...constants import Constants
from ...utils import Utils


# -------------------
## Base class for generating a msword docx
class GenBaseDocx:  # pylint: disable=too-few-public-methods

    # -------------------
    ## constructor
    def __init__(self):
        ## holds reference to the current PDF doc
        self._doc = None

    # -------------------
    ## initialize document
    #
    # @return None
    def _doc_init(self):
        self._set_page_layout()
        self._add_styles()
        self._gen_headers_footers()

    # -------------------
    ## set page layout to landscape/portrait
    #
    # @return None
    def _set_page_layout(self):
        section = self._doc.sections[0]
        self._set_section_orientation(section)

        section.header_distance = docx.shared.Inches(0.0)
        section.left_margin = docx.shared.Inches(0.5)
        section.right_margin = docx.shared.Inches(0.5)
        section.top_margin = docx.shared.Inches(0.25)
        section.bottom_margin = docx.shared.Inches(0.25)

    # -------------------
    ## add styles needed for various tables
    #
    # @return None
    def _add_styles(self):
        # Note: use keep_with_next to minimize splitting across pages
        self._add_style_table_header()
        self._add_style_table_cell()
        self._add_style_table_desc_cell()
        self._add_style_monospace()
        self._add_style_list_item()

    # -------------------
    ## add a paragraph style to the document, set common values
    #
    # @param name  the style name
    # @return the style
    def _add_style(self, name):
        style = self._doc.styles.add_style(name, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)  # pylint: disable=no-member
        style.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE  # pylint: disable=no-member
        style.paragraph_format.alignment = self._align_left()
        style.paragraph_format.keep_with_next = True
        return style

    # -------------------
    ## add a paragraph style for table headings (bold)
    #
    # @return None
    def _add_style_table_header(self):
        style = self._add_style('ver_table1_header')
        style.base_style = self._doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = docx.shared.Pt(8)
        style.font.bold = True
        style.paragraph_format.space_before = docx.shared.Inches(0.08)
        style.paragraph_format.space_after = docx.shared.Inches(0.08)

    # -------------------
    ## add a paragraph style for table cells
    #
    # @return None
    def _add_style_table_cell(self):
        style = self._add_style('ver_table1_cell')
        style.base_style = self._doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = docx.shared.Pt(8)
        style.font.bold = False
        style.paragraph_format.space_before = docx.shared.Inches(0.0)
        style.paragraph_format.space_after = docx.shared.Inches(0.0)

    # -------------------
    ## add a paragraph style for descriptions in tables
    #
    # @return None
    def _add_style_table_desc_cell(self):
        style = self._add_style('ver_table1_desc_cell')
        style.base_style = self._doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = docx.shared.Pt(8)
        style.font.bold = False
        style.paragraph_format.space_before = docx.shared.Inches(0.08)
        style.paragraph_format.space_after = docx.shared.Inches(0.0)

    # -------------------
    ## add a paragraph style with monospace font
    #
    # @return None
    def _add_style_monospace(self):
        style = self._add_style('ver_monospace')
        style.base_style = self._doc.styles['Normal']
        style.font.name = 'Courier New'
        style.font.size = docx.shared.Pt(7)
        style.font.bold = True
        style.paragraph_format.space_before = docx.shared.Inches(0.0)
        style.paragraph_format.space_after = docx.shared.Inches(0.0)

    # -------------------
    ## add a paragraph style for lists
    #
    # @return None
    def _add_style_list_item(self):
        style = self._add_style('ver_list_item')
        style.base_style = self._doc.styles['List Bullet']
        style.font.name = 'Arial'
        style.font.size = docx.shared.Pt(8)
        style.font.bold = False
        style.paragraph_format.space_before = docx.shared.Inches(0.08)
        style.paragraph_format.space_after = docx.shared.Inches(0.0)

    # -------------------
    ## add before/after space to a table
    #
    # @param size   the font size to use
    # @return None
    def _tbl_add_space(self, size):
        para = self._doc.add_paragraph()
        run = para.add_run(' ')  # must have a space!
        run.font.size = docx.shared.Pt(size)

        # 20ths of a point
        # tbl._tblPr.set(qn('w:bottomFromText'), '1440')  # pylint: disable=protected-access

        # logger.dbg(f'"{tbl._tblPr}"')  # pylint: disable=protected-access
        # tbl._tblPr.append(gap)

        # tblPr = tbl._tbl.tblPr
        # floating = OxmlElement('w:tblpPr')  # floating table properties
        # floating.set(qn('w:topFromText'), '720')  # <--------- relevant values
        # floating.set(qn('w:bottomFromText'), '720')  # <--------- relevant values
        # floating.set(qn('w:tblpY'), '721')  # <--------- relevant values
        # floating.set(qn('w:vertAnchor'), 'text')
        # tblPr.append(floating)  # floating table properties
        #
        # overlap = OxmlElement('w:tblOverlap')  # overlap properties
        # overlap.set(qn('w:val'), 'never')
        # tblPr.append(overlap)

    # -------------------
    ## set the background of the given cell to light blue
    #
    # @param cell  the cell to change
    # @return None
    def _set_cell_blue_background(self, cell):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'ADD8E6')  # light blue
        cell._tc.get_or_add_tcPr().append(shd)  # pylint: disable=protected-access

    # -------------------
    ## set the background of the given cell to light grey
    #
    # @param cell  the cell to change
    # @return None
    def _set_cell_grey_background(self, cell):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')  # light grey
        cell._tc.get_or_add_tcPr().append(shd)  # pylint: disable=protected-access

    # -------------------
    ## generate test run information
    #
    # @return None
    def _gen_test_run_details(self):
        self._doc.add_heading('Test Run Details', level=3)

        line = f"{'Test Run Type'}: {services.cfg.test_run_type}"
        self._doc.add_paragraph(line, style='List Bullet')

        line = f"{'Test Run ID'}: {services.cfg.test_run_id}"
        self._doc.add_paragraph(line, style='List Bullet')

        dts = Utils.get_dts(use_cfg_fmt=True)
        line = f"{'Document Generated'}: {dts}"
        self._doc.add_paragraph(line, style='List Bullet')

        line = f"{'medver-pytest version'}: v{Constants.version}"
        self._doc.add_paragraph(line, style='List Bullet')

    # -------------------
    ## generate title
    #
    # @param title  the title to draw
    # @return None
    def _gen_title(self, title):
        self._doc.add_heading(title, level=3)

    # -------------------
    ## generate headers and footers
    #
    # @return None
    def _gen_headers_footers(self):
        self._gen_headers()
        self._gen_footers()

    # -------------------
    ## generate headers
    #
    # @return None
    def _gen_headers(self):
        header = self._doc.sections[0].header
        if services.cfg.page_info.orientation == 'portrait':
            tbl = header.add_table(rows=1, cols=3, width=docx.shared.Inches(7.5))
        else:
            tbl = header.add_table(rows=1, cols=3, width=docx.shared.Inches(10.0))
        tbl.style = 'Normal Table'
        self._align_table_center(tbl)

        edge_width = docx.shared.Inches(1.75)
        if services.cfg.page_info.orientation == 'portrait':
            # should add to 7.5"
            tbl.columns[0].width = edge_width
            tbl.columns[1].width = docx.shared.Inches(4.00)
            tbl.columns[2].width = edge_width
        else:
            # should add to 10"
            tbl.columns[0].width = edge_width
            tbl.columns[1].width = docx.shared.Inches(6.5)
            tbl.columns[2].width = edge_width

        row = tbl.rows[0].cells
        self._gen_hf_data(row, 'header')

        para = header.add_paragraph()
        self._insert_hr(para)

    # -------------------
    ## generate footers
    #
    # @return None
    def _gen_footers(self):
        footer = self._doc.sections[0].footer
        self._insert_hr(footer.paragraphs[0])
        if services.cfg.page_info.orientation == 'portrait':
            tbl = footer.add_table(rows=1, cols=3, width=docx.shared.Inches(7.5))
        else:
            tbl = footer.add_table(rows=1, cols=3, width=docx.shared.Inches(10.0))
        tbl.style = 'Normal Table'
        self._align_table_center(tbl)
        tbl.keep_with_next = True

        edge_width = docx.shared.Inches(1.75)
        if services.cfg.page_info.orientation == 'portrait':
            # should add to 7.5"
            tbl.columns[0].width = edge_width
            tbl.columns[1].width = docx.shared.Inches(4.00)
            tbl.columns[2].width = edge_width
        else:
            # should add to 10"
            tbl.columns[0].width = edge_width
            tbl.columns[1].width = docx.shared.Inches(6.5)
            tbl.columns[2].width = edge_width

        row = tbl.rows[0].cells
        self._gen_hf_data(row, 'footer')

    # -------------------
    ## generate headers/footer data
    #
    # @param row  the row for the header/footer data
    # @param tag  indicates header or footer
    # @return None
    def _gen_hf_data(self, row, tag):
        col = 0
        if services.cfg.page_info[tag].left == '<pageno>':
            self._add_page_number(row, col, self._align_left())
        else:
            txt = services.cfg.page_info[tag].left
            txt = txt.replace('<br/>', '\n')
            boldness = False
            if txt.find('<b>') != -1 or txt.find('</b>') != -1 or txt.find('<b/>') != -1:
                boldness = True
                txt = txt.replace('<b>', '')
                txt = txt.replace('</b>', '')
                txt = txt.replace('<b/>', '')
            row[col].paragraphs[0].add_run(txt).bold = boldness
            row[col].paragraphs[0].alignment = self._align_left()

        col = 1
        if services.cfg.page_info[tag].middle == '<pageno>':
            self._add_page_number(row, col, self._align_center())
        else:
            txt = services.cfg.page_info[tag].middle
            txt = txt.replace('<br/>', '\n')
            boldness = False
            if txt.find('<b>') != -1 or txt.find('</b>') != -1 or txt.find('<b/>') != -1:
                boldness = True
                txt = txt.replace('<b>', '')
                txt = txt.replace('</b>', '')
                txt = txt.replace('<b/>', '')
            row[col].paragraphs[0].add_run(txt).bold = boldness
            row[col].paragraphs[0].alignment = self._align_center()

        col = 2
        if services.cfg.page_info[tag].right == '<pageno>':
            self._add_page_number(row, col, self._align_left())
        else:
            txt = services.cfg.page_info[tag].right
            txt = txt.replace('<br/>', '\n')
            boldness = False
            if txt.find('<b>') != -1 or txt.find('</b>') != -1 or txt.find('<b/>') != -1:
                boldness = True
                txt = txt.replace('<b>', '')
                txt = txt.replace('</b>', '')
                txt = txt.replace('<b/>', '')
            row[col].paragraphs[0].add_run(txt).bold = boldness
            row[col].paragraphs[0].alignment = self._align_left()

    # -------------------
    ## add page number to header/footer
    #
    # @param row        the cell to add the page to
    # @param col        the col number to add the page to
    # @param alignment  the paragraph alignment to use
    # @return None
    def _add_page_number(self, row, col, alignment):
        run = row[col].paragraphs[0].add_run('Page ')

        elem = OxmlElement('w:fldChar')
        elem.set(qn('w:fldCharType'), 'begin')
        run._r.append(elem)  # pylint: disable=protected-access

        elem = OxmlElement('w:instrText')
        elem.set(qn('xml:space'), 'preserve')
        elem.text = "PAGE"
        run._r.append(elem)  # pylint: disable=protected-access

        elem = OxmlElement('w:fldChar')
        elem.set(qn('w:fldCharType'), 'end')
        run._r.append(elem)  # pylint: disable=protected-access

        row[col].paragraphs[0].alignment = alignment

    # -------------------
    ## insert a header paragraph
    #
    # @param paragraph the paragraph to add to the header
    # @return None
    def _insert_hr(self, paragraph):
        para = paragraph._p  # pylint: disable=protected-access
        ppr = para.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        ppr.insert_element_before(pbdr,
                                  'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                                  'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                                  'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                                  'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                                  'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                                  'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                                  'w:pPrChange'
                                  )
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '0')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), 'auto')
        pbdr.append(bottom)

    # -------------------
    ## set alignment to left
    #
    # @return None
    def _align_left(self):
        return docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT  # pylint: disable=no-member

    # -------------------
    ## set alignment to center
    #
    # @return None
    def _align_center(self):
        return docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  # pylint: disable=no-member

    # -------------------
    ## set table alignment to center
    #
    # @param tbl  the table to align
    # @return None
    def _align_table_center(self, tbl):
        tbl.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER  # pylint: disable=no-member

    # -------------------
    ## set cell's vertical alignment to center
    #
    # @param cell the cell to align
    # @return None
    def _align_vertical_center(self, cell):
        cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER  # pylint: disable=no-member

    # -------------------
    ## set cell's vertical alignment to top
    #
    # @param cell the cell to align
    # @return None
    def _align_vertical_top(self, cell):
        cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.TOP  # pylint: disable=no-member

    # -------------------
    ## set the portrait/landscape orientation for the given section
    #
    # @param section the section to set
    # @return None
    def _set_section_orientation(self, section):
        if services.cfg.page_info.page_size == 'letter':
            # the default
            pass
        elif services.cfg.page_info.page_size == 'a4':
            section.page_height = docx.shared.Mm(297)
            section.page_width = docx.shared.Mm(210)
            section.left_margin = docx.shared.Mm(25.4)
            section.right_margin = docx.shared.Mm(25.4)
            section.top_margin = docx.shared.Mm(25.4)
            section.bottom_margin = docx.shared.Mm(25.4)
            section.header_distance = docx.shared.Mm(12.7)
            section.footer_distance = docx.shared.Mm(12.7)
        else:
            services.abort(f'unknown page_size: {services.cfg.page_info.page_size}')

        if services.cfg.page_info.orientation == 'portrait':
            section.orientation = docx.enum.section.WD_ORIENTATION.PORTRAIT  # pylint: disable=no-member
            # no rotate needed
        else:
            section.orientation = docx.enum.section.WD_ORIENTATION.LANDSCAPE  # pylint: disable=no-member
            # rotate page
            new_width = section.page_width
            new_height = section.page_height
            section.page_width = new_height
            section.page_height = new_width
