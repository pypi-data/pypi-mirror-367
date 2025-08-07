import os

from docx import Document
from docx.shared import Inches

from .gen_summary_base import GenSummaryBase
from ..common.gen_base_docx import GenBaseDocx
from ... import services
from ...constants import Constants


# -------------------
## Generates a Summary report in PDF format
class GenSummaryDocx(GenBaseDocx, GenSummaryBase):
    # -------------------
    ## constructor
    #
    # @param summary    the summary data to use
    def __init__(self, summary):
        GenBaseDocx.__init__(self)
        GenSummaryBase.__init__(self, summary)

        ## holds reference to the table being created
        self._tbl = None

    # -------------------
    ## generate the report
    #
    # @return None
    def gen(self):
        self._gen_init()

        ## see GenBaseDocx
        self._doc = Document()
        self._doc_init()

        self._gen_test_run_details()
        self._gen_title('Summary')

        self._gen_req_summary()
        self._doc.add_paragraph('')
        self._gen_proto_summary()

        self._build()

    # -------------------
    ## generate requirements summary
    #
    # @return None
    def _gen_req_summary(self):
        self._create_table()

        # gen header rows
        self._gen_header1('Requirements Summary')
        self._gen_header2()

        # total requirements
        total = self._get_total_requirements()
        self._write_row('#Requirements', total, total)

        num_failing, num_passing, num_invalid, num_missing = self._count_requirements(total, self._report_invalid)
        self._write_row('   Invalid', num_invalid, total)
        self._write_row('   PASS', num_passing, total)
        self._write_row('   FAIL', num_failing, total)
        self._write_row('   Not tested', num_missing, total)

    # -------------------
    ## callback to report any invalid requirements
    #
    # @param reqid   the reqid to report as invalid
    # @return None
    def _report_invalid(self, reqid):
        msg = f'   {reqid} not found in {services.cfg.reqmt_json_path}'
        self._write_row(msg, None, None)
        services.logger.warn(msg)

    # -------------------
    ## generate protocol summary
    #
    # @return None
    def _gen_proto_summary(self):
        self._create_table()

        # gen header rows
        self._gen_header1('Protocol Summary')
        self._gen_header2()

        total = len(self._summary['protoids'])
        self._write_row('#Protocols', total, total)

        # count failing/passing reqmts
        num_failing, num_passing = self._count_protocols()
        self._write_row('   PASS', num_passing, total)
        self._write_row('   FAIL', num_failing, total)

    # -------------------
    ## create a table with 3 columns
    #
    # @return None
    def _create_table(self):
        # TODO add 0.25 space before and after table
        self._tbl = self._doc.add_table(rows=1, cols=3)
        self._tbl.style = 'Table Grid'
        self._align_table_center(self._tbl)
        self._tbl.columns[0].width = Inches(3.20)
        self._tbl.columns[1].width = Inches(0.65)
        self._tbl.columns[2].width = Inches(0.90)

    # -------------------
    ## generate header1 row
    #
    # @param title  the text to set the header to
    # @return None
    def _gen_header1(self, title):
        row = self._tbl.rows[0].cells
        self._set_header1_cell(row[0], title)
        row[0].merge(row[1])
        row[0].merge(row[2])

    # -------------------
    ## generate header2 row
    #
    # @return None
    def _gen_header2(self):
        row = self._tbl.add_row().cells
        self._set_header2_cell(row[0], '')
        self._set_header2_cell(row[1], 'Count')
        self._set_header2_cell(row[2], 'Percentage')

    # -------------------
    ## write a row to table
    #
    # @param msg    the line header to write
    # @param count  the value to write
    # @param total  the total count
    # @return None
    def _write_row(self, msg, count, total):
        if count is None or total is None:
            pct = ''
        else:
            pct = self._pct(count, total)
        if count is None:
            count = ''

        row = self._tbl.add_row().cells
        row[0].paragraphs[0].add_run(str(msg), )
        row[0].paragraphs[0].style = 'ver_table1_cell'

        row[1].paragraphs[0].add_run(f'{count: >8}', )
        row[1].paragraphs[0].style = 'ver_table1_cell'

        row[2].paragraphs[0].add_run(f'{pct: >8}', )
        row[2].paragraphs[0].style = 'ver_table1_cell'

    # -------------------
    ## set header1 cell to the given text
    #
    # @param cell  the cell to set
    # @param text  the text to set the cell to
    # @return None
    def _set_header1_cell(self, cell, text):
        cell.paragraphs[0].add_run(text)
        cell.paragraphs[0].style = 'ver_table1_header'
        self._set_cell_blue_background(cell)
        self._align_vertical_top(cell)

    # -------------------
    ## set header2 cell to the given text
    #
    # @param cell  the cell to set
    # @param text  the text to set the cell to
    # @return None
    def _set_header2_cell(self, cell, text):
        cell.paragraphs[0].add_run(text)
        cell.paragraphs[0].style = 'ver_table1_header'
        self._set_cell_grey_background(cell)
        self._align_vertical_top(cell)

    # -------------------
    ## build the document
    #
    # @return None
    def _build(self):
        path = os.path.join(services.cfg.outdir, f'{Constants.summary_fname}.docx')
        self._doc.save(path)
