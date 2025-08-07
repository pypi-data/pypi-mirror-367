import docx
import docx.enum.table
import docx.enum.text
import docx.shared
from docx import Document

from .gen_tp_base import GenTpBase
from ..common.gen_base_docx import GenBaseDocx
from ... import services
from ...result_summary import ResultSummary


# -------------------
## Generates a Test Protocol report in MSWord docx format
class GenTpDocx(GenBaseDocx, GenTpBase):
    # -------------------
    ## constructor
    #
    # @param protocols   the data to use
    # @param do_results  generate results or not
    def __init__(self, protocols, do_results=True):
        ## see base class doc
        self._testers = None
        ## see base clas doc
        self._doc_path = None

        GenBaseDocx.__init__(self)
        GenTpBase.__init__(self, protocols, do_results)

        self._init('docx', protocols, do_results)

    # -------------------
    ## generate the report
    #
    # @return None
    def gen(self):
        self._gen_init()
        # uncomment to debug:
        # services.logger.dbg(f'protocols: {json.dumps(self._protocols, indent=2)}')

        ## see GenBaseDocx
        self._doc = Document()
        self._doc_init()

        self._gen_test_run_details()
        title = self._get_title()
        self._gen_title(title)

        self._doc.add_paragraph('')
        for _, protocol in self._protocols.items():
            self._gen_protocol(protocol)

        self._build()

    # -------------------
    ## generate a protocol table using the given data
    #
    # @param protocol   the data to use
    # @return None
    def _gen_protocol(self, protocol):
        tbl = self._gen_protocol_table()
        self._gen_protocol_info(tbl, protocol)

        requirements = []
        need_space = self._gen_steps_info(tbl, protocol['steps'], requirements)
        self._update_protocol_table(tbl, need_space)

        self._tbl_add_space(1)
        self._gen_reqmt_table(requirements)

        if self._testers:
            self._tbl_add_space(1)
            self._gen_testers()

        self._tbl_add_space(2)

    # -------------------
    ## generate protocol info
    #
    # @param tbl        the table to write to
    # @param protocol   the data to use
    # @return None
    def _gen_protocol_info(self, tbl, protocol):
        self._gen_protocol_info_row0(tbl, protocol)
        self._gen_protocol_info_row1(tbl, protocol)
        self._gen_protocol_info_row2(tbl, protocol)
        self._gen_protocol_info_row3(tbl, protocol)
        self._gen_protocol_info_row4(tbl, protocol)

    # -------------------
    ## generate protocol info for row 0 of the header section
    # contains protocol id and description, start date
    #
    # @param tbl        the table to write to
    # @param protocol   the data to use
    # @return None
    def _gen_protocol_info_row0(self, tbl, protocol):
        row_cells = tbl.add_row().cells

        if self._do_results:
            start_date = protocol['start_date']
        else:
            start_date = ''

        self._gen_protocol_info_item(row_cells[0], row_cells[3], protocol['proto_id'], protocol['desc'])
        self._gen_protocol_info_item(row_cells[4], row_cells[6], 'Start date', start_date)

    # -------------------
    ## generate protocol info for row 1 of the header section
    # contains location of the protocol
    #
    # @param tbl        the table to write to
    # @param protocol   the data to use
    # @return None
    def _gen_protocol_info_row1(self, tbl, protocol):
        row_cells = tbl.add_row().cells

        who = f'{protocol["executed_by"]} at {protocol["location"]}'
        self._gen_protocol_info_item(row_cells[0], row_cells[3], 'At', who)
        self._gen_protocol_info_item(row_cells[4], row_cells[6], 'Requirements', 'see below')

    # -------------------
    ## generate protocol info for row 2 of the header section
    # contains softawre version and serial number of DUT
    #
    # @param tbl        the table to write to
    # @param protocol   the data to use
    # @return None
    def _gen_protocol_info_row2(self, tbl, protocol):
        row_cells = tbl.add_row().cells

        self._gen_protocol_info_item(row_cells[0], row_cells[3], 'Software Version', protocol['dut_version'])
        self._gen_protocol_info_item(row_cells[4], row_cells[6], 'Serial number', protocol['dut_serialno'])

    # -------------------
    ## generate protocol info for row 3 of the header section
    # contains Objectives of the protocol
    #
    # @param tbl        the table to write to
    # @param protocol   the data to use
    # @return None
    def _gen_protocol_info_row3(self, tbl, protocol):
        row_cells = tbl.add_row().cells
        self._gen_protocol_info_list(row_cells[0], row_cells[6], 'Objectives', protocol['objectives'])

    # -------------------
    ## generate protocol info for row 4 of the header section
    # contains Preconditions and Deviations if any
    #
    # @param tbl        the table to write to
    # @param protocol   the data to use
    # @return None
    def _gen_protocol_info_row4(self, tbl, protocol):
        row_cells = tbl.add_row().cells
        self._gen_protocol_info_list(row_cells[0], row_cells[3], 'Preconditions', protocol['preconditions'])
        self._gen_protocol_info_list(row_cells[4], row_cells[6], 'Deviations', protocol['deviations'])

    # -------------------
    ## generate protocol info for the given cell
    #
    # @param cell       to write to
    # @param last_cell  the last cell of the column span
    # @param title      the title of the cell
    # @param item       the data to use
    # @return None
    def _gen_protocol_info_item(self, cell, last_cell, title, item):
        self._cell_init(cell)
        self._cell_add_bold_text(cell, f'{title}: ')
        self._cell_add_text(cell, item)
        cell.merge(last_cell)
        self._set_cell_blue_background(cell)

    # -------------------
    ## generate protocol info for the given cell in list format
    #
    # @param cell       to write to
    # @param last_cell  the last cell of the column span
    # @param tag        the tag for the list
    # @param items      the list of items to use
    # @return None
    def _gen_protocol_info_list(self, cell, last_cell, tag, items):
        self._gen_list(cell, tag, items)
        cell.merge(last_cell)
        self._set_cell_blue_background(cell)

    # -------------------
    ## generate a list using the given items
    #
    # @param cell     the cell to write the list to
    # @param tag      the name of the list
    # @param items    the list of data to use
    # @return a paragraph containing the formatted list
    def _gen_list(self, cell, tag, items):
        self._cell_init(cell)
        self._cell_add_bold_text(cell, f'{tag}: ')

        if len(items) == 0:
            list_items = 'N/A'
            self._cell_add_text(cell, list_items)
        elif len(items) == 1:
            list_items = items[0]
            self._cell_add_text(cell, list_items)
        else:
            for item in items:
                self._cell_add_listitem(cell, item)

    # -------------------
    ## generate title row and one row per step
    #
    # @param tbl            the table to write to
    # @param steps          the steps data
    # @param requirements   the requirements found in all steps
    # @return None
    def _gen_steps_info(self, tbl, steps, requirements):
        # row 5: title row for step info
        row_cells = tbl.add_row().cells
        self._set_table_title(row_cells[0], 'Step')
        self._set_table_title(row_cells[1], 'Req.')
        self._set_table_title(row_cells[2], 'Execution Instructions')
        self._set_table_title(row_cells[3], 'Expected')
        self._set_table_title(row_cells[4], 'Actual')
        self._set_table_title(row_cells[5], 'Pass/\nFail')
        self._set_table_title(row_cells[6], 'Info')

        need_space = False
        if len(steps) == 0:
            row_cells = tbl.add_row().cells
            self._set_table_cell(row_cells[0], 'N/A')
            self._set_table_cell(row_cells[2], 'No steps found')
        else:
            # row N: one row for each step
            stepno = 0
            for step in steps:
                stepno += 1
                ns = self._gen_step(tbl, stepno, step, requirements)
                need_space = need_space or ns

        return need_space

    # -------------------
    ## generate a step in the test protocol
    #
    # @param tbl            the table to generate into
    # @param stepno         the step number of this step
    # @param step           the step description
    # @param requirements   the one or more reqids for this step
    # @return a paragraph containing the formatted list
    def _gen_step(self, tbl, stepno, step, requirements):
        # default is a passed, and empty result
        rs = ResultSummary()
        self._get_overall_result(step, rs)

        row_cells = tbl.add_row().cells
        self._set_table_cell(row_cells[0], self._get_stepno('docx', stepno))
        row_cells[0].paragraphs[0].alignment = self._align_center()
        self._set_table_cell(row_cells[1], self._get_reqids('docx', rs, requirements))
        self._set_table_cell(row_cells[2], self._get_desc('docx', step))
        self._set_table_cell_mono(row_cells[3], self._get_expected('docx', rs))
        self._set_table_cell_mono(row_cells[4], self._get_actual('docx', rs))
        self._set_table_cell(row_cells[5], self._get_result('docx', rs))
        row_cells[5].paragraphs[0].alignment = self._align_center()
        self._set_table_cell(row_cells[6], self._get_details('docx', rs, step))

        if self._do_results:
            self._gen_step_comments(tbl, step)

        # if the actual or expected is long-ish, indicate those columsn will need more space
        if services.cfg.page_info.orientation == 'portrait':
            max_len = 10
        else:
            max_len = 20
        if len(str(rs.actual_formatted)) > max_len:
            return True
        if len(str(rs.expected_formatted)) > max_len:
            return True

        return False

    # -------------------
    ## generate the comments for the given step
    #
    # @param tbl    the table to add the comments to
    # @param step   the step data
    # @return None
    def _gen_step_comments(self, tbl, step):
        for comment in step['comments']:
            para = f'Note: {comment}'
            row_cells = tbl.add_row().cells
            self._set_table_cell(row_cells[2], para)

    # -------------------
    ## generate a protocol table
    #
    # @return tbl object
    def _gen_protocol_table(self):
        tbl = self._doc.add_table(rows=0, cols=7)
        tbl.style = 'Table Grid'
        self._align_table_center(tbl)

        return tbl

    # -------------------
    ## update protocol table layout
    #
    # @param tbl          the table to write to
    # @param need_space   indicates if the table requires a different format because of space layout
    # @return None
    def _update_protocol_table(self, tbl, need_space):
        # these widths are the same portrait, landscape, needs space or not
        stepno_width = docx.shared.Inches(0.41)
        reqmt_width = docx.shared.Inches(0.65)
        details_width = docx.shared.Inches(1.47)
        result_width = docx.shared.Inches(0.47)
        if services.cfg.page_info.orientation == 'portrait':
            # portrait: (7.5" width)
            if need_space:
                value_width = docx.shared.Inches(1.31)
                desc_width = docx.shared.Inches(1.88)
            else:
                value_width = docx.shared.Inches(0.70)
                desc_width = docx.shared.Inches(3.10)
        else:
            # landscape (10" width)
            if need_space:
                value_width = docx.shared.Inches(1.68)
                desc_width = docx.shared.Inches(3.64)
            else:
                value_width = docx.shared.Inches(1.00)
                desc_width = docx.shared.Inches(5.00)

        tbl.columns[0].width = stepno_width  # stepno
        tbl.columns[1].width = reqmt_width  # reqmt.
        tbl.columns[2].width = desc_width  # desc
        tbl.columns[3].width = value_width  # expected
        tbl.columns[4].width = value_width  # actual
        tbl.columns[5].width = result_width  # result
        tbl.columns[6].width = details_width  # details

        # uncomment to debug
        # total = 0.0
        # units = 913780.48780487  # width units per inch
        # for i in range(0, 6 + 1):
        #     w = tbl.columns[i].width
        #     total += w
        #     logger.dbg(f"    {i} {w / units: >8.2f} docx.shared.Inches")
        # # there are 913780 per inch
        # logger.dbg(f"{total / units: >8.2f} docx.shared.Inches, need_space:{need_space}")

    # -------------------
    ## generate a requirement table
    #
    # @param requirements   a list of reqmt info
    # @return None
    def _gen_reqmt_table(self, requirements):
        # row 0: title row
        tbl = self._doc.add_table(rows=0, cols=2)
        tbl.style = 'Table Grid'
        self._align_table_center(tbl)
        #  = tbl._tblPr

        if services.cfg.page_info.orientation == 'portrait':
            # portrait: 7.5"
            tbl.columns[0].width = docx.shared.Inches(0.85)  # reqmt.
            tbl.columns[1].width = docx.shared.Inches(6.65)  # desc
        else:
            # landscape: 10"
            tbl.columns[0].width = docx.shared.Inches(0.85)  # reqmt.
            tbl.columns[1].width = docx.shared.Inches(9.15)  # desc

        row_cells = tbl.add_row().cells
        self._set_table_title(row_cells[0], 'Req.')
        self._set_table_title(row_cells[1], 'Desc.')

        if len(requirements) == 0:
            row_cells = tbl.add_row().cells
            self._set_table_cell(row_cells[0], 'N/A')
            self._set_table_cell(row_cells[1], 'No requirements found')
        else:
            for reqmt in sorted(requirements):
                row_cells = tbl.add_row().cells
                self._set_table_cell(row_cells[0], reqmt)

                if reqmt in self._requirements:
                    desc = self._requirements[reqmt]['desc']
                else:
                    desc = f'Could not find {reqmt} in file: {services.cfg.reqmt_json_path}'
                self._set_table_cell(row_cells[1], desc)

    # -------------------
    ## generate a signature area
    #
    # @return None
    def _gen_testers(self):
        self._gen_title('Manual Tester Signatures')

        # row 0: title row
        tbl = self._doc.add_table(rows=0, cols=3)
        tbl.style = 'Table Grid'
        self._align_table_center(tbl)

        if services.cfg.page_info.orientation == 'portrait':
            # portrait: 7.5"
            tbl.columns[0].width = docx.shared.Inches(0.50)  # initials
            tbl.columns[1].width = docx.shared.Inches(3.50)  # name
            tbl.columns[2].width = docx.shared.Inches(3.50)  # signature
        else:
            # landscape: 10"
            tbl.columns[0].width = docx.shared.Inches(0.50)  # reqmt.
            tbl.columns[1].width = docx.shared.Inches(3.25)  # name
            tbl.columns[2].width = docx.shared.Inches(6.25)  # signature

        row_cells = tbl.add_row().cells
        self._set_table_title(row_cells[0], '')
        self._set_table_title(row_cells[1], 'Tester')
        self._set_table_title(row_cells[2], 'Signature')

        for initials in sorted(self._testers):
            if initials == 'manual':
                initials = ''
                name = ''
            elif initials in services.cfg.testers:
                name = services.cfg.testers[initials]
            else:
                name = ''
            row_cells = tbl.add_row().cells
            self._set_table_cell(row_cells[0], initials)
            self._set_table_cell(row_cells[1], name)
            self._set_table_cell(row_cells[2], '\n\n')

    # -------------------
    ## set a table title to the given text
    #
    # @param cell   the table cell to write to
    # @param text   the title text to write
    # @return None
    def _set_table_title(self, cell, text):
        cell.text = ''
        cell.paragraphs[0].add_run(text).bold = True
        cell.paragraphs[0].style = 'ver_table1_header'
        self._set_cell_grey_background(cell)
        self._align_vertical_center(cell)

    # -------------------
    ## set a table cell to the given text
    #
    # @param cell   the table cell to write to
    # @param text   the text to write
    # @return None
    def _set_table_cell(self, cell, text):
        cell.text = ''
        cell.paragraphs[0].add_run(text)
        cell.paragraphs[0].style = 'ver_table1_cell'
        self._align_vertical_center(cell)

    # -------------------
    ## set a table cell to the given text in monospace font
    #
    # @param cell   the table cell to write to
    # @param text   the text to write
    # @return None
    def _set_table_cell_mono(self, cell, text):
        cell.text = ''
        cell.paragraphs[0].add_run(text)
        cell.paragraphs[0].style = 'ver_monospace'
        self._align_vertical_center(cell)

    # -------------------
    ## initialize a table cell
    #
    # @param cell   the table cell to write to
    # @return None
    def _cell_init(self, cell):
        cell.text = ''
        cell.paragraphs[0].style = 'ver_table1_cell'
        self._align_vertical_center(cell)

    # -------------------
    ## add bold text to the given table cell
    #
    # @param cell   the table cell to write to
    # @param text   the text to write
    # @return None
    def _cell_add_bold_text(self, cell, text):
        run = cell.paragraphs[0].add_run(text)
        run.bold = True

    # -------------------
    ## add text to the given table cell
    #
    # @param cell   the table cell to write to
    # @param text   the text to write
    # @return None
    def _cell_add_text(self, cell, text):
        cell.paragraphs[0].add_run(text)

    # -------------------
    ## add a list item the given table cell
    #
    # @param cell   the table cell to write to
    # @param item   the list item to write
    # @return None
    def _cell_add_listitem(self, cell, item):
        para = cell.add_paragraph(style='ver_list_item')
        para.add_run(item)

    # -------------------
    ## build the document with the current list of elements
    #
    # @return None
    def _build(self):
        self._doc.save(self._doc_path)
