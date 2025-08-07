#!/usr/bin/env python3
"""
DOC/DOCX图片提取器命令行工具
"""

import argparse
import json
import logging
import sys
from pathlib import Path
from typing import List, Optional

from ..core.extractor import extract_images, to_ascii_dirname
from ..core.config import Config, load_config


def setup_logging(level: str = "INFO") -> None:
    """设置日志"""
    logging.basicConfig(
        level=getattr(logging, level.upper()),
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )


def extract_command(args) -> None:
    """提取图片命令"""
    config = load_config(args.config) if args.config else Config()
    
    # 设置日志级别
    setup_logging(args.verbose)
    
    # 设置输出目录
    base_dir = args.output or config.base_image_dir
    
    results = []
    for doc_path in args.files:
        print(f"正在处理: {doc_path}")
        result = extract_images(doc_path, base_dir)
        results.append({
            'file': doc_path,
            'result': result
        })
        
        if result['success']:
            print(f"  ✓ 成功提取 {result['count']} 张图片到: {result['output_dir']}")
        else:
            print(f"  ✗ 失败: {result['msg']}")
    
    # 输出JSON结果
    if args.json:
        print(json.dumps(results, ensure_ascii=False, indent=2))


def preview_command(args) -> None:
    """预览DOC/DOCX结构命令"""
    import zipfile
    import olefile
    
    for doc_path in args.files:
        print(f"\n=== {doc_path} ===")
        
        try:
            if doc_path.lower().endswith('.docx'):
                # 使用python-docx预览DOCX文档结构
                try:
                    from docx import Document
                    doc = Document(doc_path)
                    print(f"段落数量: {len(doc.paragraphs)}")
                    print(f"表格数量: {len(doc.tables)}")
                except ImportError:
                    print("未安装python-docx，跳过文档结构分析")
                
                # 预览ZIP结构
                with zipfile.ZipFile(doc_path, 'r') as zf:
                    media_files = [f for f in zf.namelist() if f.startswith('word/media/')]
                    print(f"媒体文件数量: {len(media_files)}")
                    
                    if args.verbose == "DEBUG":
                        print("媒体文件列表:")
                        for media_file in media_files:
                            print(f"  - {media_file}")
                            
            elif doc_path.lower().endswith('.doc'):
                # 预览DOC文件结构
                if olefile.isOleFile(doc_path):
                    with olefile.OleFileIO(doc_path) as ole:
                        streams = ole.listdir()
                        print(f"OLE流数量: {len(streams)}")
                        
                        # 查找可能包含图片的流
                        image_streams = []
                        for stream in streams:
                            stream_name = '/'.join(stream)
                            if any(keyword in stream_name.lower() for keyword in ['data', 'objinfo', 'object', '1table', 'worddocument']):
                                image_streams.append(stream_name)
                        
                        print(f"可能包含图片的流: {len(image_streams)}")
                        
                        if args.verbose == "DEBUG":
                            print("所有流:")
                            for stream in streams:
                                print(f"  - {'/'.join(stream)}")
                else:
                    print("  ✗ 不是有效的DOC文件")
            else:
                print("  ✗ 不支持的文件格式，仅支持.doc和.docx")
                        
        except Exception as e:
            print(f"  ✗ 错误: {e}")


def convert_command(args) -> None:
    """文件名转换命令"""
    for filename in args.filenames:
        ascii_name = to_ascii_dirname(filename)
        print(f"{filename} -> {ascii_name}")


def config_command(args) -> None:
    """配置管理命令"""
    if args.action == "show":
        config = load_config(args.config) if args.config else Config()
        print(json.dumps(config._config, ensure_ascii=False, indent=2))
    
    elif args.action == "create":
        config_path = args.output or "config.json"
        config = Config()
        config.save_to_file(config_path)
        print(f"配置文件已创建: {config_path}")


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description="DOC/DOCX图片提取器命令行工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例用法:
  # 提取单个文件的图片
  python cli.py extract document.docx
  python cli.py extract document.doc
  
  # 提取多个文件的图片到指定目录
  python cli.py extract -o images/ doc1.docx doc2.doc
  
  # 预览文档结构
  python cli.py preview document.docx
  python cli.py preview document.doc
  
  # 转换文件名为ASCII
  python cli.py convert "测试文档.docx" "另一个文档.doc"
  
  # 显示当前配置
  python cli.py config show
        """
    )
    
    parser.add_argument(
        "-v", "--verbose",
        choices=["DEBUG", "INFO", "WARNING", "ERROR"],
        default="INFO",
        help="日志级别"
    )
    
    parser.add_argument(
        "-c", "--config",
        help="配置文件路径"
    )
    
    subparsers = parser.add_subparsers(dest="command", help="可用命令")
    
    # 提取命令
    extract_parser = subparsers.add_parser("extract", help="提取DOC/DOCX文件中的图片")
    extract_parser.add_argument("files", nargs="+", help="DOC/DOCX文件路径")
    extract_parser.add_argument("-o", "--output", help="输出目录")
    extract_parser.add_argument("--json", action="store_true", help="输出JSON格式结果")
    extract_parser.set_defaults(func=extract_command)
    
    # 预览命令
    preview_parser = subparsers.add_parser("preview", help="预览DOC/DOCX文件结构")
    preview_parser.add_argument("files", nargs="+", help="DOC/DOCX文件路径")
    preview_parser.set_defaults(func=preview_command)
    
    # 转换命令
    convert_parser = subparsers.add_parser("convert", help="转换文件名为ASCII")
    convert_parser.add_argument("filenames", nargs="+", help="要转换的文件名")
    convert_parser.set_defaults(func=convert_command)
    
    # 配置命令
    config_parser = subparsers.add_parser("config", help="配置管理")
    config_parser.add_argument(
        "action",
        choices=["show", "create"],
        help="配置操作: show(显示配置), create(创建配置文件)"
    )
    config_parser.add_argument("-o", "--output", help="输出文件路径(仅用于create)")
    config_parser.set_defaults(func=config_command)
    
    args = parser.parse_args()
    
    if not args.command:
        parser.print_help()
        sys.exit(1)
    
    try:
        args.func(args)
    except KeyboardInterrupt:
        print("\n操作已取消")
        sys.exit(1)
    except Exception as e:
        print(f"错误: {e}")
        if args.verbose == "DEBUG":
            import traceback
            traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()