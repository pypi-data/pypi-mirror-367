"""Main module."""

import pandas as pd
from matplotlib import pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import argparse
from datetime import datetime
import os


class ReportGenerator:
    def __init__(self, excel_file='budget.xlsx', output_file=None):
        self.excel_file = excel_file
        self.document = None
        self.data = None
        self._content_sections = None  # Cache for content.md sections

        # Create reports directory if it doesn't exist
        reports_dir = Path('reports')
        reports_dir.mkdir(exist_ok=True)

        # Simplified naming: project_report_yyyy-mm-dd.docx in reports folder
        if output_file is None:
            base_name = f'project_report_{datetime.now().strftime("%Y-%m-%d")}.docx'
            base_path = reports_dir / base_name
            self.output_file = str(self._get_unique_filename(base_path))
        else:
            # If custom filename provided, still put it in reports folder
            self.output_file = str(reports_dir / output_file)

    def _get_unique_filename(self, base_filepath):
        """Ensure unique filename by adding increment if file exists"""
        if not base_filepath.exists():
            return base_filepath

        # If file exists, add increment: project_report_2025-01-25_v2.docx
        name_stem = base_filepath.stem  # project_report_2025-01-25
        extension = base_filepath.suffix  # .docx
        parent_dir = base_filepath.parent

        counter = 2
        while True:
            new_filename = f"{name_stem}_v{counter}{extension}"
            new_filepath = parent_dir / new_filename
            if not new_filepath.exists():
                print(f"📝 File exists, creating new version: {new_filename}")
                return new_filepath
            counter += 1

            # Safety break to avoid infinite loop
            if counter > 100:
                # Use timestamp as fallback
                timestamp = datetime.now().strftime("%H%M%S")
                return parent_dir / f"{name_stem}_{timestamp}{extension}"

    def _load_content_sections(self):
        """Load and parse content.md file into sections"""
        if self._content_sections is not None:
            return self._content_sections

        content_file = 'content.md'
        self._content_sections = {}

        if not os.path.exists(content_file):
            print(f"⚠️  Content file '{content_file}' not found")
            return self._content_sections

        try:
            with open(content_file, 'r', encoding='utf-8') as f:
                content = f.read()

            # Split content by headers (## Header Name)
            import re
            sections = re.split(r'^## (.+)$', content, flags=re.MULTILINE)

            # First section before any header is ignored
            for i in range(1, len(sections), 2):
                if i + 1 < len(sections):
                    header_name = sections[i].strip()
                    section_content = sections[i + 1].strip()

                    # Normalize header names (remove spaces, make lowercase)
                    normalized_name = header_name.lower().replace(' ', '_').replace('-', '_')
                    self._content_sections[normalized_name] = {
                        'original_name': header_name,
                        'content': section_content
                    }

            print(
                f"📄 Loaded {len(self._content_sections)} sections from content.md")
            return self._content_sections

        except (FileNotFoundError, UnicodeDecodeError) as e:
            print(f"❌ Error loading content.md: {e}")
            return {}

    def read_section_from_content(self, section_name):
        """Read a specific section from content.md"""
        sections = self._load_content_sections()
        normalized_name = section_name.lower().replace(' ', '_').replace('-', '_')

        if normalized_name in sections:
            return sections[normalized_name]['content']
        else:
            available_sections = list(sections.keys())
            print(f"⚠️  Section '{section_name}' not found in content.md")
            if available_sections:
                print(
                    f"   Available sections: {', '.join(available_sections)}")
            return None

    def add_markdown_content(self, section_name, default_content=None):
        """Add markdown content from content.md to the document"""
        content = self.read_section_from_content(section_name)

        if not content and default_content:
            content = default_content
            print(f"⚠️  Using default content for section '{section_name}'")

        if content:
            # Split content into lines for processing
            lines = content.split('\n')

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Handle bullet points (markdown style) - removed bold formatting
                if line.startswith('- '):
                    bullet_text = line[2:].strip()
                    self.document.add_paragraph(
                        bullet_text, style='List Bullet')
                else:
                    # Regular paragraph
                    self.document.add_paragraph(line)
        else:
            # Add placeholder if no content found
            self.document.add_paragraph(
                f"[Content for {section_name} section]")

    def add_full_markdown_file(self, markdown_file_path, start_header_level=1):
        """Add complete markdown file to document with automatic header detection.

        Args:
            markdown_file_path (str): Path to the markdown file
            start_header_level (int): Starting header level for Word document (1-9)

        Returns:
            bool: True if successful, False otherwise
        """
        try:
            from .gen_auto import add_markdown_to_existing_document
        except ImportError:
            try:
                from autorpt.gen_auto import add_markdown_to_existing_document
            except ImportError:
                import sys
                current_dir = Path(__file__).parent
                sys.path.insert(0, str(current_dir))
                try:
                    from gen_auto import add_markdown_to_existing_document
                except ImportError:
                    print("❌ Error: Could not import auto-generation module.")
                    return False

        return add_markdown_to_existing_document(self.document, markdown_file_path, start_header_level)

    def add_excel_table(self, excel_file_path, sheet_name=None, table_title=None, start_header_level=1):
        """Add Excel table to document with automatic formatting.

        Args:
            excel_file_path (str): Path to the Excel file
            sheet_name (str): Optional sheet name to read
            table_title (str): Optional title for the table
            start_header_level (int): Starting header level for Word document (1-9)

        Returns:
            bool: True if successful, False otherwise
        """
        try:
            from .gen_auto import add_excel_table_to_existing_document
        except ImportError:
            try:
                from autorpt.gen_auto import add_excel_table_to_existing_document
            except ImportError:
                import sys
                current_dir = Path(__file__).parent
                sys.path.insert(0, str(current_dir))
                try:
                    from gen_auto import add_excel_table_to_existing_document
                except ImportError:
                    print("❌ Error: Could not import auto-generation module.")
                    return False

        return add_excel_table_to_existing_document(
            self.document, excel_file_path, sheet_name, table_title, start_header_level
        )

    def add_mixed_content(self, content_files, start_header_level=1):
        """Add multiple markdown and Excel files to the document.

        Args:
            content_files (list): List of file paths or file dictionaries
            start_header_level (int): Starting header level for Word document (1-9)

        Returns:
            dict: Results summary with success/failure counts
        """
        try:
            from .gen_auto import add_mixed_content_to_existing_document
        except ImportError:
            try:
                from autorpt.gen_auto import add_mixed_content_to_existing_document
            except ImportError:
                import sys
                current_dir = Path(__file__).parent
                sys.path.insert(0, str(current_dir))
                try:
                    from gen_auto import add_mixed_content_to_existing_document
                except ImportError:
                    print("❌ Error: Could not import auto-generation module.")
                    return {'success': 0, 'failed': 1, 'files': []}

        return add_mixed_content_to_existing_document(self.document, content_files, start_header_level)

    def add_all_content_from_folder(self, content_folder="reports", start_header_level=1):
        """Automatically discover and add all markdown and Excel files from a folder.

        Args:
            content_folder (str): Folder to scan for content files
            start_header_level (int): Starting header level for Word document (1-9)

        Returns:
            dict: Results summary with success/failure counts
        """
        try:
            from .gen_auto import AutoReportGenerator
        except ImportError:
            try:
                from autorpt.gen_auto import AutoReportGenerator
            except ImportError:
                import sys
                current_dir = Path(__file__).parent
                sys.path.insert(0, str(current_dir))
                try:
                    from gen_auto import AutoReportGenerator
                except ImportError:
                    print("❌ Error: Could not import auto-generation module.")
                    return {'success': 0, 'failed': 1, 'files': [], 'discovered': 0}

        # Use the existing document
        temp_generator = AutoReportGenerator(self.document, content_folder)
        return temp_generator.add_all_content_from_folder(start_header_level)

    def _remove_table_borders(self, table):
        """Remove all borders from table for cleaner look"""
        from docx.oxml import parse_xml

        # Create XML for no borders
        no_border_xml = """
        <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="none"/>
            <w:left w:val="none"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="none"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
        """

        borders_element = parse_xml(no_border_xml)
        # Access protected member for low-level XML manipulation (required for border removal)
        table._tbl.tblPr.append(borders_element)  # type: ignore[attr-defined]

    def _format_cell_alignment(self, cell, column_index, is_header=False, is_total_row=False):
        """Format cell alignment based on content type"""

        # Set alignment based on column content
        if column_index == 0:  # First column (category names)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:  # Numeric columns
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Make header row bold
        if is_header or is_total_row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    if is_total_row:
                        run.font.size = Pt(11)

    def load_data(self):
        """Load Excel data"""
        try:
            # Read Excel file - handle multiple sheets if needed
            excel_file = pd.ExcelFile(self.excel_file)

            # Use first sheet by default or specific sheet if provided
            sheet_names = excel_file.sheet_names
            print(
                f"📊 Found {len(sheet_names)} sheet(s): {', '.join(sheet_names)}")

            # Load the first sheet (or could be made configurable)
            self.data = pd.read_excel(
                self.excel_file, sheet_name=sheet_names[0])
            print(
                f"✅ Successfully loaded {len(self.data)} rows from {sheet_names[0]}")

            return True

        except (FileNotFoundError, ValueError, pd.errors.EmptyDataError) as e:
            print(f"❌ Error loading Excel file: {e}")
            return False

    def create_budget_table(self):
        """Create formatted budget table"""
        if self.data is None:
            return

        # Add Budget Table Header
        header = self.document.add_heading('Budget Overview', level=2)
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Create table with data
        table = self.document.add_table(
            rows=len(self.data) + 1, cols=len(self.data.columns))

        # Remove borders for cleaner look
        self._remove_table_borders(table)

        # Add header row
        header_cells = table.rows[0].cells
        for i, column in enumerate(self.data.columns):
            header_cells[i].text = str(column)
            self._format_cell_alignment(header_cells[i], i, is_header=True)

        # Add data rows
        for i, (_, row) in enumerate(self.data.iterrows(), start=1):
            data_cells = table.rows[i].cells
            for j, value in enumerate(row):
                # Format numbers as currency if they appear to be monetary
                if isinstance(value, (int, float)) and j > 0:
                    data_cells[j].text = f"${value:,.2f}"
                else:
                    data_cells[j].text = str(value)

                # Check if this is a total row (last row or contains "Total" in first column)
                is_total_row = (i == len(self.data)) or (
                    "total" in str(row.iloc[0]).lower())
                self._format_cell_alignment(
                    data_cells[j], j, is_total_row=is_total_row)

    def add_chart_section(self):
        """Add budget chart section"""
        # Add Chart section header
        chart_header = self.document.add_heading(
            'Budget Visualization', level=2)
        chart_header.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Create and save chart
        self.create_budget_chart()

        # Check if chart was created successfully
        chart_path = 'budget_chart.png'
        if os.path.exists(chart_path):
            # Add chart to document
            chart_paragraph = self.document.add_paragraph()
            chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = chart_paragraph.runs[0] if chart_paragraph.runs else chart_paragraph.add_run(
            )
            run.add_picture(chart_path, width=Inches(6))

            print(f"✅ Chart added to document: {chart_path}")
        else:
            print("❌ Chart file not found, skipping chart section")

    def create_budget_chart(self):
        """Create budget visualization chart"""
        if self.data is None:
            return

        try:
            # Prepare data for chart - assume first column is categories, others are values
            categories = self.data.iloc[:, 0].tolist()

            # If there are multiple value columns, use the last one or sum them
            if len(self.data.columns) > 2:
                # Sum numeric columns for total budget
                numeric_cols = self.data.select_dtypes(
                    include=['number']).columns
                values = self.data[numeric_cols].sum(axis=1).tolist()
            else:
                # Use second column as values
                values = self.data.iloc[:, 1].tolist()

            # Create chart
            plt.figure(figsize=(10, 6))
            bars = plt.bar(categories, values)

            # Customize chart
            plt.title('Budget Overview', fontsize=16,
                      fontweight='bold', pad=20)
            plt.xlabel('Categories', fontsize=12)
            plt.ylabel('Amount ($)', fontsize=12)

            # Format y-axis as currency
            plt.gca().yaxis.set_major_formatter(
                plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))

            # Rotate x-axis labels if too long
            plt.xticks(rotation=45, ha='right')

            # Add value labels on bars
            for bar, value in zip(bars, values):
                height = bar.get_height()
                plt.text(bar.get_x() + bar.get_width()/2., height,
                         f'${value:,.0f}', ha='center', va='bottom', fontsize=10)

            # Adjust layout and save
            plt.tight_layout()
            plt.savefig('budget_chart.png', dpi=300, bbox_inches='tight')
            plt.close()

            print("✅ Budget chart created successfully")

        except (ValueError, OSError) as e:
            print(f"❌ Error creating budget chart: {e}")

    def add_summary_section(self):
        """Add executive summary section"""
        summary_header = self.document.add_heading(
            'Executive Summary', level=2)
        summary_header.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Default summary content
        default_intro = """This report provides a comprehensive overview of the project budget and financial analysis. The budget has been carefully structured to ensure optimal resource allocation while maintaining project objectives."""

        self.add_markdown_content('Summary', default_intro)

    def add_methodology_section(self):
        """Add methodology section"""
        method_header = self.document.add_heading('Methodology', level=2)
        method_header.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Default methodology content
        default_methodology = """The budget analysis methodology includes:
- Comprehensive review of all cost categories
- Analysis of resource requirements and allocation
- Risk assessment and contingency planning
- Alignment with project objectives and timelines"""

        self.add_markdown_content(
            'Methodology', default_methodology)

    def add_findings_section(self):
        """Add key findings section with data insights"""
        findings_header = self.document.add_heading('Key Findings', level=2)
        findings_header.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Generate insights from data
        if self.data is not None:
            try:
                # Calculate basic statistics
                numeric_cols = self.data.select_dtypes(
                    include=['number']).columns
                if len(numeric_cols) > 0:
                    total_budget = self.data[numeric_cols].sum().sum()
                    max_category = self.data.iloc[:, 0][self.data[numeric_cols].sum(
                        axis=1).idxmax()]
                    min_category = self.data.iloc[:, 0][self.data[numeric_cols].sum(
                        axis=1).idxmin()]

                    insights = f"""Based on the budget analysis, the following key findings have been identified:

- Total project budget: ${total_budget:,.2f}
- Highest budget allocation: {max_category}
- Lowest budget allocation: {min_category}
- Number of budget categories: {len(self.data)}

The budget distribution shows a strategic allocation of resources across different project components."""

                    self.add_markdown_content('Findings', insights)
                else:
                    # Fallback if no numeric data
                    default_findings = """Key findings from the budget analysis include strategic resource allocation and alignment with project objectives."""
                    self.add_markdown_content('Findings', default_findings)

            except (ValueError, KeyError, TypeError) as e:
                print(f"⚠️  Error generating insights: {e}")
                default_findings = """Key findings from the budget analysis include strategic resource allocation and alignment with project objectives."""
                self.add_markdown_content('Findings', default_findings)
        else:
            default_key_points = """Key findings include:
- Strategic budget allocation across project components
- Alignment with organizational objectives
- Comprehensive resource planning"""

            self.add_markdown_content('key_points', default_key_points)

    def add_chart_description(self):
        """Add chart description section"""
        # Default chart description
        chart_desc = """The budget visualization chart provides a clear overview of resource allocation across different project categories. This visual representation helps stakeholders understand the distribution of funds and identify key investment areas."""

        self.add_markdown_content('chart_description', chart_desc)

    def add_recommendations_section(self):
        """Add recommendations section"""
        rec_header = self.document.add_heading('Recommendations', level=2)
        rec_header.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Try to get recommendations from content.md, with fallback
        recommendations_content = self.read_section_from_content(
            'Recommendations')

        if not recommendations_content:
            # Default recommendations
            default_recommendations = """Based on the budget analysis, the following recommendations are proposed:

- Continue monitoring budget performance against established benchmarks
- Implement regular review cycles to ensure optimal resource utilization
- Consider potential cost optimization opportunities
- Maintain flexibility for adjustments based on project evolution
- Establish clear reporting mechanisms for budget tracking"""

            # Split into paragraphs and add to document
            for paragraph in default_recommendations.strip().split('\n\n'):
                if paragraph.strip():
                    lines = paragraph.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line.startswith('- '):
                            # Bullet point
                            bullet_text = line[2:].strip()
                            self.document.add_paragraph(
                                bullet_text, style='List Bullet')
                        elif line:
                            # Regular paragraph
                            self.document.add_paragraph(line)
        else:
            # Use content from content.md
            self.add_markdown_content('Recommendations')

    def add_conclusion_section(self):
        """Add conclusion section"""
        conclusion_header = self.document.add_heading('Conclusion', level=2)
        conclusion_header.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Default conclusion content
        default_conclusion = """This budget analysis provides a comprehensive foundation for informed decision-making. The structured approach to resource allocation supports project success while maintaining fiscal responsibility. Regular monitoring and adjustment will ensure optimal outcomes."""

        self.add_markdown_content('Conclusion', default_conclusion)

    def add_challenges_section(self):
        """Add challenges section"""
        challenges_header = self.document.add_heading(
            'Challenges and Mitigation', level=2)
        challenges_header.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Default challenges content
        default_challenges = """Potential challenges and mitigation strategies include:
- Budget variance management through regular monitoring
- Resource availability through strategic planning
- Market fluctuations through contingency planning"""

        self.add_markdown_content('Challenges', default_challenges)

    def add_next_steps_section(self):
        """Add next steps section"""
        next_steps_header = self.document.add_heading('Next Steps', level=2)
        next_steps_header.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Default next steps content
        default_next_steps = """Recommended next steps include:
- Implementation of budget monitoring systems
- Regular stakeholder communication and reporting
- Periodic review and adjustment processes
- Development of contingency planning protocols"""

        self.add_markdown_content(
            'Next Steps', default_next_steps)

    def rpt_pdf(self, word_file=None):
        """Convert Word report to PDF

        Args:
            word_file (str): Path to Word document. If None, uses self.output_file

        Returns:
            str: Path to created PDF file if successful, None if failed
        """
        # Import pdf module - handle both package and script execution contexts
        try:
            # Try relative import first (for package installation)
            from .pdf import convert_to_pdf as pdf_convert_func
        except ImportError:
            try:
                # Try absolute import (for installed package)
                from autorpt.pdf import convert_to_pdf as pdf_convert_func
            except ImportError:
                # If both fail, try adding current directory to path (for script execution)
                import sys
                current_dir = Path(__file__).parent
                sys.path.insert(0, str(current_dir))
                try:
                    from pdf import convert_to_pdf as pdf_convert_func
                except ImportError:
                    print(
                        "❌ Error: Could not import PDF conversion module. Make sure docx2pdf is installed.")
                    return None

        # Use provided file or default output file
        source_file = word_file or self.output_file

        if not source_file or not os.path.exists(source_file):
            print(f"❌ Word document not found: {source_file}")
            return None

        success, result = pdf_convert_func(source_file)

        if success:
            print(f"🎉 PDF conversion completed: {result}")
            return result
        else:
            print(f"❌ PDF conversion failed: {result}")
            return None

    def generate_report(self):
        """Generate complete report"""
        try:
            print(f"📊 Loading data from {self.excel_file}...")
            if not self.load_data():
                return False

            print("📄 Creating Word document...")
            self.document = Document()

            # Add title
            title = self.document.add_heading('Project Budget Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add date
            date_para = self.document.add_paragraph(
                f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add sections
            print("📝 Adding report sections...")
            self.add_summary_section()
            self.add_methodology_section()
            self.create_budget_table()
            self.add_findings_section()
            self.add_chart_section()
            self.add_chart_description()
            self.add_recommendations_section()
            self.add_challenges_section()
            self.add_next_steps_section()
            self.add_conclusion_section()

            # Save document
            self.document.save(self.output_file)
            print(f"✅ Report generated successfully: {self.output_file}")
            return True

        except (OSError, PermissionError) as e:
            print(f"❌ Error generating report: {e}")
            return False

    def save_document(self, filename=None):
        """Save the document to file"""
        if self.document is None:
            print("❌ No document to save")
            return False

        save_path = filename or self.output_file
        try:
            self.document.save(save_path)
            print(f"💾 Document saved: {save_path}")
            return True
        except (OSError, PermissionError) as e:
            print(f"❌ Error saving document: {e}")
            return False


# For backward compatibility
AutoRpt = ReportGenerator


def main():
    """Main function with command line argument support"""
    parser = argparse.ArgumentParser(
        description='Generate automated budget report',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  autorpt                                      # Show usage help
  autorpt --generate                           # Generate Word report
  autorpt --pdf                                # Convert latest report to PDF
  autorpt --pdf-all                            # Convert all reports to PDF
  autorpt --generate --pdf                     # Generate Word + PDF
  autorpt --generate -i budget.xlsx --pdf     # Custom input with PDF
  autorpt --generate -m report.md              # Add markdown file with auto headers
  autorpt --generate -m report.md --pdf       # Add markdown + generate PDF
  autorpt --generate --excel data.xlsx        # Add Excel table
  autorpt --generate --mixed file1.md data.xlsx    # Add multiple files
  autorpt --generate --auto-content           # Auto-discover files in reports folder
  autorpt --generate --auto-content --content-folder "data"  # Auto-discover in custom folder
        """)

    parser.add_argument('--input', '-i', default='budget.xlsx',
                        help='Input Excel file (default: budget.xlsx)')
    parser.add_argument('--output', '-o',
                        help='Output Word document filename')
    parser.add_argument('--markdown', '-m',
                        help='Markdown file to include with auto-detected headers')
    parser.add_argument('--excel', '-e',
                        help='Excel file to add as table')
    parser.add_argument('--sheet',
                        help='Specific sheet name in Excel file (for --excel option)')
    parser.add_argument('--table-title',
                        help='Title for the Excel table (for --excel option)')
    parser.add_argument('--mixed', nargs='+',
                        help='Add multiple markdown and/or Excel files')
    parser.add_argument('--auto-content', action='store_true',
                        help='Automatically discover and add all .md and .xlsx files from content folder')
    parser.add_argument('--content-folder', default='reports',
                        help='Folder to scan for auto-content files (default: reports)')
    parser.add_argument('--pdf', '-p', action='store_true',
                        help='Convert most recent report to PDF (primary use case)')
    parser.add_argument('--pdf-only', action='store_true',
                        help='(Legacy) Same as --pdf - convert existing reports to PDF')
    parser.add_argument('--pdf-all', action='store_true',
                        help='Convert all Word reports in reports/ folder to PDF')
    parser.add_argument('--generate', action='store_true',
                        help='Generate new report (combine with --pdf for generate+convert)')
    parser.add_argument('--verbose', '-v', action='store_true',
                        help='Enable verbose output')

    args = parser.parse_args()

    if args.verbose:
        print("🔧 Verbose mode enabled")
        print(f"📁 Input file: {args.input}")
        print(f"📄 Output file: {args.output or 'auto-generated'}")
        if args.pdf:
            print("📄 PDF conversion enabled")
        if args.pdf_only:
            print("📄 PDF-only mode (no report generation)")
        if args.pdf_all:
            print("📄 Converting all reports to PDF")

    # Handle PDF-only operations (including the new simple --pdf behavior)
    if args.pdf_only or args.pdf_all or (args.pdf and not args.generate):
        # Import pdf module - handle both package and script execution contexts
        try:
            # Try relative import first (for package installation)
            from .pdf import convert_to_pdf as pdf_convert_func, convert_all_reports
        except ImportError:
            try:
                # Try absolute import (for installed package)
                from autorpt.pdf import convert_to_pdf as pdf_convert_func, convert_all_reports
            except ImportError:
                # If both fail, try adding current directory to path (for script execution)
                import sys
                current_dir = Path(__file__).parent
                sys.path.insert(0, str(current_dir))
                try:
                    from pdf import convert_to_pdf as pdf_convert_func, convert_all_reports
                except ImportError:
                    print(
                        "❌ Error: Could not import PDF conversion module. Make sure docx2pdf is installed.")
                    return 1

        if args.pdf_all:
            print("📁 Converting all Word reports to PDF...")
            results = convert_all_reports("reports")
            success = results["failed"] == 0
            if success:
                print(
                    f"🎉 Successfully converted {results['success']} report(s) to PDF!")
            else:
                print(
                    f"⚠️  Converted {results['success']} report(s), but {results['failed']} failed")
        elif args.pdf_only or (args.pdf and not args.generate):
            # Convert the most recent report or specified output file
            if args.output:
                report_file = f"reports/{args.output}"
            else:
                # Find the most recent report
                reports_dir = Path('reports')
                if reports_dir.exists():
                    docx_files = sorted(reports_dir.glob(
                        "*.docx"), key=lambda x: x.stat().st_mtime, reverse=True)
                    if docx_files:
                        report_file = str(docx_files[0])
                        print(
                            f"🔍 Converting most recent report: {docx_files[0].name}")
                    else:
                        print("❌ No Word reports found in reports/ folder")
                        return 1
                else:
                    print("❌ Reports directory not found")
                    return 1

            success, result = pdf_convert_func(report_file)
            if success:
                print("🎉 PDF conversion completed successfully!")
            else:
                print(f"❌ PDF conversion failed: {result}")

        return 0 if success else 1

    # Normal report generation flow (only when --generate is specified or no PDF-only flags)
    if args.generate or not (args.pdf or args.pdf_only or args.pdf_all):
        generator = ReportGenerator(args.input, args.output)

        # Generate basic report first
        print("🚀 Starting report generation...")
        success = generator.generate_report()

        if not success:
            print("❌ Failed to generate base report")
            return 1

        # Add additional content if specified
        if args.markdown:
            if os.path.exists(args.markdown):
                print(f"📄 Adding markdown file: {args.markdown}")
                markdown_success = generator.add_full_markdown_file(
                    args.markdown, start_header_level=1)
                if markdown_success:
                    print("✅ Markdown content added successfully")
                    generator.save_document()
                else:
                    print(
                        "⚠️  Failed to add markdown content, but report generated successfully")
            else:
                print(f"❌ Markdown file not found: {args.markdown}")

        if args.excel:
            if os.path.exists(args.excel):
                print(f"📊 Adding Excel table: {args.excel}")
                excel_success = generator.add_excel_table(
                    args.excel, args.sheet, args.table_title, start_header_level=1)
                if excel_success:
                    print("✅ Excel table added successfully")
                    generator.save_document()
                else:
                    print(
                        "⚠️  Failed to add Excel table, but report generated successfully")
            else:
                print(f"❌ Excel file not found: {args.excel}")

        if args.mixed:
            print(f"📁 Adding mixed content: {', '.join(args.mixed)}")
            existing_files = [f for f in args.mixed if os.path.exists(f)]
            if existing_files:
                results = generator.add_mixed_content(
                    existing_files, start_header_level=1)
                print(f"✅ Successfully added {results['success']} files")
                if results['failed'] > 0:
                    print(f"⚠️  Failed to add {results['failed']} files")
                generator.save_document()
            else:
                print("❌ None of the specified mixed content files were found")

        # Auto-discover and add content from folder
        if args.auto_content:
            print(
                f"🔍 Auto-discovering content files in '{args.content_folder}' folder...")
            results = generator.add_all_content_from_folder(
                args.content_folder, start_header_level=1)
            if results['discovered'] > 0:
                print(
                    f"✅ Successfully processed {results['success']}/{results['discovered']} discovered files")
                if results['failed'] > 0:
                    print(f"⚠️  Failed to process {results['failed']} files")
                generator.save_document()
            else:
                print(
                    f"📂 No additional content files found in '{args.content_folder}' folder")

        # Convert to PDF if requested during generation
        if args.pdf and args.generate:
            pdf_success = generator.rpt_pdf()
            if not pdf_success:
                print("⚠️  Report generated successfully but PDF conversion failed")

        return 0 if success else 1

    # If we get here, it means no generation was requested but no PDF-only operation either
    # This handles the case where someone runs 'autorpt' with no arguments
    else:
        print("ℹ️  No operation specified. Use --generate to create a new report, --pdf to convert latest report, or --pdf-all to convert all reports.")
        return 0


def convert_docx_to_pdf(word_file):
    """Standalone function to convert a Word document to PDF

    Args:
        word_file (str): Path to the Word document to convert

    Returns:
        bool: True if conversion successful, False otherwise
    """
    # Import pdf module - handle both package and script execution contexts
    try:
        # Try relative import first (for package installation)
        from .pdf import convert_to_pdf as pdf_convert
    except ImportError:
        try:
            # Try absolute import (for installed package)
            from autorpt.pdf import convert_to_pdf as pdf_convert
        except ImportError:
            # If both fail, try adding current directory to path (for script execution)
            import sys
            current_dir = Path(__file__).parent
            sys.path.insert(0, str(current_dir))
            try:
                from pdf import convert_to_pdf as pdf_convert
            except ImportError:
                print(
                    "❌ Error: Could not import PDF conversion module. Make sure docx2pdf is installed.")
                return False

    success, _ = pdf_convert(word_file)
    return success


if __name__ == "__main__":
    exit(main())
