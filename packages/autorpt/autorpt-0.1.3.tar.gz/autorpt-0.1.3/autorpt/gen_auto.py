"""General auto-generation module for converting markdown and Excel to Word documents."""

import re
import pandas as pd
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


class AutoReportGenerator:
    """Convert markdown files and Excel tables to Word document sections with proper formatting."""

    def __init__(self, document=None, content_folder="reports"):
        """Initialize the generator with an optional Word document.

        Args:
            document: Existing Word document object, or None to create new one
            content_folder (str): Folder to scan for additional content files
        """
        self.document = document if document else Document()
        self.content_folder = Path(content_folder)

    def discover_content_files(self, include_patterns=None, exclude_patterns=None):
        """Automatically discover markdown and Excel files in the content folder.

        Args:
            include_patterns (list): List of glob patterns to include (default: ['*.md', '*.xlsx', '*.xls'])
            exclude_patterns (list): List of patterns to exclude (default: ['*report*.docx'])

        Returns:
            dict: Dictionary with 'markdown' and 'excel' file lists
        """
        if include_patterns is None:
            include_patterns = ['*.md', '*.markdown', '*.xlsx', '*.xls']

        if exclude_patterns is None:
            exclude_patterns = ['*report*.docx', '*report*.pdf']

        discovered = {'markdown': [], 'excel': [], 'all': []}

        if not self.content_folder.exists():
            print(f"📁 Content folder '{self.content_folder}' not found")
            return discovered

        try:
            # Find all matching files
            all_files = []
            for pattern in include_patterns:
                all_files.extend(self.content_folder.glob(pattern))

            # Filter out excluded patterns
            filtered_files = []
            for file_path in all_files:
                exclude_file = False
                for exclude_pattern in exclude_patterns:
                    if file_path.match(exclude_pattern):
                        exclude_file = True
                        break
                if not exclude_file:
                    filtered_files.append(file_path)

            # Categorize files
            for file_path in filtered_files:
                if file_path.suffix.lower() in ['.md', '.markdown']:
                    discovered['markdown'].append(file_path)
                elif file_path.suffix.lower() in ['.xlsx', '.xls']:
                    discovered['excel'].append(file_path)
                discovered['all'].append(file_path)

            # Sort files by name for consistent ordering
            discovered['markdown'].sort(key=lambda x: x.name)
            discovered['excel'].sort(key=lambda x: x.name)
            discovered['all'].sort(key=lambda x: x.name)

            if discovered['all']:
                print(
                    f"📂 Discovered {len(discovered['all'])} content files in '{self.content_folder}':")
                for file_path in discovered['all']:
                    print(f"   📄 {file_path.name}")
            else:
                print(
                    f"📂 No additional content files found in '{self.content_folder}'")

        except (OSError, PermissionError) as e:
            print(f"❌ Error discovering content files: {e}")

        return discovered

    def parse_markdown_file(self, markdown_file_path):
        """Parse a markdown file and extract structured content.

        Args:
            markdown_file_path (str or Path): Path to the markdown file

        Returns:
            list: List of content blocks with headers and content
        """
        markdown_path = Path(markdown_file_path)

        if not markdown_path.exists():
            print(f"Warning: Markdown file {markdown_path} not found")
            return []

        try:
            with open(markdown_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except (FileNotFoundError, PermissionError, UnicodeDecodeError) as e:
            print(f"Error reading {markdown_path}: {e}")
            return []

        return self._parse_markdown_content(content)

    def parse_excel_file(self, excel_file_path, sheet_name=None, table_title=None):
        """Parse an Excel file and extract table data.

        Args:
            excel_file_path (str or Path): Path to the Excel file
            sheet_name (str): Optional sheet name to read (default: first sheet)
            table_title (str): Optional title for the table

        Returns:
            dict: Dictionary with table data and metadata
        """
        excel_path = Path(excel_file_path)

        if not excel_path.exists():
            print(f"Warning: Excel file {excel_path} not found")
            return None

        try:
            # Read Excel file
            if sheet_name:
                data = pd.read_excel(excel_path, sheet_name=sheet_name)
            else:
                data = pd.read_excel(excel_path)

            # Get sheet info for title if not provided
            if not table_title:
                if sheet_name:
                    table_title = f"Table: {sheet_name}"
                else:
                    table_title = f"Table: {excel_path.stem}"

            return {
                'data': data,
                'title': table_title,
                'file_path': excel_path,
                'sheet_name': sheet_name,
                'columns': list(data.columns),
                'rows': len(data)
            }

        except (FileNotFoundError, PermissionError, ValueError, ImportError) as e:
            print(f"Error reading Excel file {excel_path}: {e}")
            return None

    def _parse_markdown_content(self, content):
        """Parse markdown content into structured blocks.

        Args:
            content (str): Raw markdown content

        Returns:
            list: List of dictionaries with 'type', 'level', 'text', and 'content'
        """
        blocks = []
        lines = content.split('\n')
        current_section = None
        current_content = []

        for line in lines:
            # Check for headers (# ## ### etc.)
            header_match = re.match(r'^(#{1,6})\s+(.+)', line.strip())

            if header_match:
                # Save previous section if exists
                if current_section is not None:
                    current_section['content'] = '\n'.join(
                        current_content).strip()
                    blocks.append(current_section)

                # Start new section
                level = len(header_match.group(1))  # Number of # characters
                title = header_match.group(2).strip()

                current_section = {
                    'type': 'header',
                    'level': level,
                    'text': title,
                    'content': ''
                }
                current_content = []
            else:
                # Add line to current content
                current_content.append(line)

        # Save last section
        if current_section:
            current_section['content'] = '\n'.join(current_content).strip()
            blocks.append(current_section)

        # If no headers found, treat entire content as one block
        if not blocks and content.strip():
            blocks.append({
                'type': 'content',
                'level': 0,
                'text': '',
                'content': content.strip()
            })

        return blocks

    def add_markdown_to_document(self, markdown_file_path, start_header_level=1):
        """Add markdown content to the Word document with proper formatting.

        Args:
            markdown_file_path (str or Path): Path to the markdown file
            start_header_level (int): Starting header level for Word document (1-9)

        Returns:
            bool: True if successful, False otherwise
        """
        blocks = self.parse_markdown_file(markdown_file_path)

        if not blocks:
            return False

        for block in blocks:
            if block['type'] == 'header':
                # Add header to document
                word_level = min(start_header_level + block['level'] - 1, 9)
                self.document.add_heading(block['text'], level=word_level)

                # Add content if any
                if block['content'].strip():
                    self._add_formatted_content(block['content'])
            else:
                # Add content without header
                if block['content'].strip():
                    self._add_formatted_content(block['content'])

        return True

    def add_excel_table_to_document(self, excel_file_path, sheet_name=None, table_title=None,
                                    include_header=True, start_header_level=1):
        """Add Excel table to the Word document with proper formatting.

        Args:
            excel_file_path (str or Path): Path to the Excel file
            sheet_name (str): Optional sheet name to read
            table_title (str): Optional title for the table
            include_header (bool): Whether to add a header before the table
            start_header_level (int): Header level for the table title

        Returns:
            bool: True if successful, False otherwise
        """
        excel_data = self.parse_excel_file(
            excel_file_path, sheet_name, table_title)

        if not excel_data:
            return False

        data = excel_data['data']
        title = excel_data['title']

        # Add header if requested
        if include_header:
            self.document.add_heading(title, level=start_header_level)

        # Add summary paragraph
        summary_text = f"Data summary: {excel_data['rows']} rows, {len(excel_data['columns'])} columns"
        self.document.add_paragraph(summary_text)

        # Create table
        table = self.document.add_table(rows=1, cols=len(data.columns))
        table.style = 'Table Grid'  # Use a clean table style

        # Add header row
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(data.columns):
            hdr_cells[i].text = str(column_name)
            # Make header bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for _, row_data in data.iterrows():
            row_cells = table.add_row().cells

            for j, cell_value in enumerate(row_data):
                # Format cell value
                if pd.api.types.is_numeric_dtype(type(cell_value)) and pd.notna(cell_value):
                    # Format numbers with appropriate precision
                    if isinstance(cell_value, float):
                        if cell_value == int(cell_value):
                            row_cells[j].text = f"{int(cell_value):,}"
                        else:
                            row_cells[j].text = f"{cell_value:,.2f}"
                    else:
                        row_cells[j].text = f"{cell_value:,}"
                else:
                    row_cells[j].text = str(
                        cell_value) if pd.notna(cell_value) else ""

                # Right-align numeric columns
                if pd.api.types.is_numeric_dtype(type(cell_value)):
                    for paragraph in row_cells[j].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        print(f"✅ Added Excel table: {title}")
        return True

    def add_mixed_content_to_document(self, content_files, start_header_level=1):
        """Add multiple markdown and Excel files to the document.

        Args:
            content_files (list): List of file paths or dictionaries with file info
            start_header_level (int): Starting header level

        Returns:
            dict: Results summary with success/failure counts
        """
        results = {'success': 0, 'failed': 0, 'files': []}

        for item in content_files:
            if isinstance(item, str):
                file_path = Path(item)
                options = {}
            else:
                file_path = Path(item['file'])
                options = item.get('options', {})

            try:
                if file_path.suffix.lower() in ['.md', '.markdown']:
                    # Handle markdown file
                    success = self.add_markdown_to_document(
                        file_path,
                        options.get('start_header_level', start_header_level)
                    )
                elif file_path.suffix.lower() in ['.xlsx', '.xls']:
                    # Handle Excel file
                    success = self.add_excel_table_to_document(
                        file_path,
                        sheet_name=options.get('sheet_name'),
                        table_title=options.get('table_title'),
                        include_header=options.get('include_header', True),
                        start_header_level=options.get(
                            'start_header_level', start_header_level)
                    )
                else:
                    print(f"⚠️  Unsupported file type: {file_path.suffix}")
                    success = False

                if success:
                    results['success'] += 1
                    results['files'].append(
                        {'file': str(file_path), 'status': 'success'})
                else:
                    results['failed'] += 1
                    results['files'].append(
                        {'file': str(file_path), 'status': 'failed'})

            except (FileNotFoundError, PermissionError, ValueError) as e:
                print(f"❌ Error processing {file_path}: {e}")
                results['failed'] += 1
                results['files'].append(
                    {'file': str(file_path), 'status': 'error', 'error': str(e)})

        return results

    def add_all_content_from_folder(self, start_header_level=1, include_patterns=None, exclude_patterns=None):
        """Automatically discover and add all content files from the content folder.

        Args:
            start_header_level (int): Starting header level
            include_patterns (list): List of glob patterns to include
            exclude_patterns (list): List of patterns to exclude

        Returns:
            dict: Results summary with success/failure counts
        """
        discovered = self.discover_content_files(
            include_patterns, exclude_patterns)

        if not discovered['all']:
            return {'success': 0, 'failed': 0, 'files': [], 'discovered': 0}

        # Convert Path objects to strings for add_mixed_content_to_document
        file_paths = [str(file_path) for file_path in discovered['all']]

        results = self.add_mixed_content_to_document(
            file_paths, start_header_level)
        results['discovered'] = len(discovered['all'])

        return results

    def _add_formatted_content(self, content):
        """Add formatted content to the document with markdown-style formatting.

        Args:
            content (str): Content to add with markdown formatting
        """
        # Split content into paragraphs
        paragraphs = content.split('\n\n')

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Handle different types of content
            if para.startswith('- ') or para.startswith('* '):
                self._add_bullet_list(para)
            elif para.startswith('1. ') or re.match(r'^\d+\.\s', para):
                self._add_numbered_list(para)
            elif para.startswith('> '):
                self._add_blockquote(para)
            else:
                self._add_regular_paragraph(para)

    def _add_bullet_list(self, content):
        """Add bullet list items to the document."""
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('- ') or line.startswith('* '):
                bullet_text = line[2:].strip()
                if bullet_text:
                    self.document.add_paragraph(
                        bullet_text, style='List Bullet')

    def _add_numbered_list(self, content):
        """Add numbered list items to the document."""
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if re.match(r'^\d+\.\s', line):
                # Remove the number and period
                numbered_text = re.sub(r'^\d+\.\s', '', line)
                if numbered_text:
                    self.document.add_paragraph(
                        numbered_text, style='List Number')

    def _add_blockquote(self, content):
        """Add blockquote content to the document."""
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('> '):
                quote_text = line[2:].strip()
                if quote_text:
                    para = self.document.add_paragraph(quote_text)
                    # Style as italic for blockquotes
                    for run in para.runs:
                        run.italic = True

    def _add_regular_paragraph(self, content):
        """Add regular paragraph with basic markdown formatting."""
        # Handle basic markdown formatting
        content = self._apply_basic_formatting(content)

        # Split by lines for multi-line paragraphs
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line:
                self.document.add_paragraph(line)

    def _apply_basic_formatting(self, text):
        """Apply basic markdown formatting (bold, italic) to text.

        Note: This is a simplified implementation. Word formatting
        would require more complex handling of runs.

        Args:
            text (str): Text with markdown formatting

        Returns:
            str: Text with formatting markers removed (for now)
        """
        # For now, just remove markdown formatting markers
        # In a more advanced implementation, we'd convert these to Word formatting
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)      # Italic
        text = re.sub(r'`(.*?)`', r'\1', text)        # Code

        return text

    def generate_auto_report(self, content_files, output_path, document_title=None):
        """Generate a complete Word document from multiple content files.

        Args:
            content_files (list): List of markdown and Excel files to include
            output_path (str or Path): Output path for the Word document
            document_title (str): Optional title for the document

        Returns:
            bool: True if successful, False otherwise
        """
        # Create new document
        self.document = Document()

        # Add document title if provided
        if document_title:
            self.document.add_heading(document_title, level=0)

        # Process all content files
        results = self.add_mixed_content_to_document(content_files)

        # Save document
        try:
            self.document.save(output_path)
            print(f"✅ Auto-generated report saved: {output_path}")
            print(
                f"📊 Processed {results['success']} files successfully, {results['failed']} failed")
            return True
        except (OSError, PermissionError) as e:
            print(f"❌ Error saving document: {e}")
            return False


# Convenience functions for easy use
def convert_mixed_content_to_word(content_files, output_file, document_title=None):
    """Convert multiple markdown and Excel files to a Word document.

    Args:
        content_files (list): List of file paths or file dictionaries
        output_file (str): Output path for Word document
        document_title (str): Optional document title

    Returns:
        bool: True if successful, False otherwise
    """
    generator = AutoReportGenerator()
    return generator.generate_auto_report(content_files, output_file, document_title)


def auto_generate_from_folder(content_folder="reports", output_file=None, document_title=None,
                              include_patterns=None, exclude_patterns=None):
    """Automatically discover and convert all content files in a folder to a Word document.

    Args:
        content_folder (str): Folder to scan for content files (default: "reports")
        output_file (str): Output path for Word document (auto-generated if None)
        document_title (str): Optional document title
        include_patterns (list): File patterns to include
        exclude_patterns (list): File patterns to exclude

    Returns:
        dict: Results with success/failure counts and output file path
    """
    generator = AutoReportGenerator(content_folder=content_folder)

    # Generate output filename if not provided
    if not output_file:
        from datetime import datetime
        content_path = Path(content_folder)
        timestamp = datetime.now().strftime("%Y-%m-%d_%H%M")
        output_file = content_path / f"auto_report_{timestamp}.docx"

    # Create new document with title
    generator.document = Document()
    if document_title:
        generator.document.add_heading(document_title, level=0)
    elif not document_title:
        generator.document.add_heading("Auto-Generated Report", level=0)
        # Add generation info
        from datetime import datetime
        date_para = generator.document.add_paragraph(
            f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
        )
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process all discovered files
    results = generator.add_all_content_from_folder(
        start_header_level=1,
        include_patterns=include_patterns,
        exclude_patterns=exclude_patterns
    )

    # Save document
    try:
        generator.document.save(output_file)
        results['output_file'] = str(output_file)
        results['saved'] = True
        print(f"✅ Auto-generated report saved: {output_file}")
        print(
            f"📊 Processed {results['success']}/{results['discovered']} files successfully")
        if results['failed'] > 0:
            print(f"⚠️  {results['failed']} files failed to process")
    except (OSError, PermissionError) as e:
        results['saved'] = False
        results['error'] = str(e)
        print(f"❌ Error saving document: {e}")

    return results


def add_mixed_content_to_existing_document(document, content_files, start_level=1):
    """Add markdown and Excel content to an existing Word document.

    Args:
        document: Word document object
        content_files (list): List of content files to add
        start_level (int): Starting header level (1-9)

    Returns:
        dict: Results summary
    """
    generator = AutoReportGenerator(document)
    return generator.add_mixed_content_to_document(content_files, start_level)


def add_markdown_to_existing_document(document, markdown_file, start_level=1):
    """Add markdown content to an existing Word document.

    Args:
        document: Word document object
        markdown_file (str): Path to markdown file
        start_level (int): Starting header level (1-9)

    Returns:
        bool: True if successful, False otherwise
    """
    generator = AutoReportGenerator(document)
    return generator.add_markdown_to_document(markdown_file, start_level)


def add_excel_table_to_existing_document(document, excel_file, sheet_name=None,
                                         table_title=None, start_level=1):
    """Add Excel table to an existing Word document.

    Args:
        document: Word document object
        excel_file (str): Path to Excel file
        sheet_name (str): Optional sheet name
        table_title (str): Optional table title
        start_level (int): Starting header level (1-9)

    Returns:
        bool: True if successful, False otherwise
    """
    generator = AutoReportGenerator(document)
    return generator.add_excel_table_to_document(
        excel_file, sheet_name, table_title, True, start_level
    )
