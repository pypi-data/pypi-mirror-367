#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
SAST漏洞修复服务 - MCP UVX 兼容版本
提供docx文档的SAST报告解析、漏洞修复建议生成和报告导出功能

Requirements:
- mcp>=1.0.0
- python-docx>=1.1.0

Usage with uvx:
uvx --from . sast-fixer-mcp
"""

import csv
import os
import tempfile
import logging
import asyncio
import json
import re
from typing import Dict, Any, List

from mcp.server.models import InitializationOptions
import mcp.types as types
from mcp.server import NotificationOptions, Server
import mcp.server.stdio

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("请运行: pip install python-docx")
    exit(1)

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler(os.path.join(tempfile.gettempdir(), "sast_fixer_mcp.log")),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("SASTFixerMCPServer")

# 创建MCP服务器
server = Server("sast-fixer-mcp")

def get_paragraph_heading_level(paragraph):
    """获取段落标题级别"""
    if paragraph.style.name.startswith('Heading'):
        try:
            return int(paragraph.style.name.replace('Heading', ''))
        except ValueError:
            return 0
    return 0

def process_table(table):
    """处理表格数据"""
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_text = ""
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    cell_text += paragraph.text + " "
            row_data.append(cell_text.strip())
        table_data.append(row_data)
    return table_data

def get_docx_elements_in_order(document):
    """按顺序获取文档元素"""
    body = document._element.body
    elements = []
    for child in body.iterchildren():
        if child.tag.endswith('p'):
            for paragraph in document.paragraphs:
                if paragraph._element is child:
                    elements.append(('paragraph', paragraph))
                    break
        elif child.tag.endswith('tbl'):
            for table in document.tables:
                if table._element is child:
                    elements.append(('table', table))
                    break
    return elements

def parse_docx_to_json(docx_path):
    """解析docx文档为JSON结构"""
    try:
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"文件不存在: {docx_path}")
        
        doc = Document(docx_path)
        elements = get_docx_elements_in_order(doc)
        result = []
        heading_stack = [(-1, None, result)]
        current_content = []

        for elem_type, elem in elements:
            if elem_type == 'paragraph':
                level = get_paragraph_heading_level(elem)
                
                if level > 0:
                    if heading_stack[-1][1] is not None:
                        heading_stack[-1][1]["content"] = current_content
                    
                    heading_node = {
                        "title": elem.text,
                        "content": []
                    }
                    current_content = []
                    
                    while heading_stack[-1][0] >= level:
                        heading_stack.pop()
                    
                    parent_level, parent_node, parent_container = heading_stack[-1]
                    
                    if isinstance(parent_container, list):
                        parent_container.append(heading_node)
                    else:
                        if "children" not in parent_container:
                            parent_container["children"] = []
                        parent_container["children"].append(heading_node)
                    
                    heading_stack.append((level, heading_node, heading_node))
                elif elem.text.strip():
                    current_content.append({"type": "text", "value": elem.text})
            
            elif elem_type == 'table':
                table_data = process_table(elem)
                current_content.append({"type": "table", "value": table_data})
        
        if heading_stack[-1][1] is not None:
            heading_stack[-1][1]["content"] = current_content
        
        return result
    except Exception as e:
        logger.error(f"解析docx文件时出错: {e}")
        raise

def transform_json(json_data):
    """转换JSON结构为漏洞报告格式"""
    result = {}
    issue_level_pattern = re.compile(r"【(低危|中危|高危|提示)】.*?漏洞数：(\d+)")

    for item in json_data:
        main_title = item["title"]
        result[main_title] = []
        if "children" not in item:
            continue
            
        for child in item["children"]:
            match = issue_level_pattern.search(child["title"])
            if match:
                issue_level = match.group(1)
                issue_count = match.group(2)

                # 映射漏洞等级
                level_mapping = {
                    "提示": "Notice",
                    "低危": "Low", 
                    "中危": "Medium",
                    "高危": "High"
                }
                issue_level = level_mapping.get(issue_level, issue_level)
                    
                # 只处理中危和高危漏洞
                if issue_level not in ["Medium", "High"]:
                    continue

                issue = {
                    "issue_title": child["title"],
                    "issue_level": issue_level,
                    "issue_count": issue_count,
                    "issue_desc": "",
                    "fix_advice": "",
                    "code_sample": "",
                    "code_list": []
                }
                
                if "children" not in child:
                    continue
                    
                for section in child["children"]:
                    section_title = section["title"]
                    section_content = section.get("content", [])
                    
                    if section_title == "漏洞描述":
                        for content_item in section_content:
                            if content_item["type"] == "text":
                                issue["issue_desc"] = content_item["value"]
                    
                    elif section_title == "修复建议":
                        for content_item in section_content:
                            if content_item["type"] == "text":
                                issue["fix_advice"] = content_item["value"]
                    
                    elif section_title == "代码示例":
                        for content_item in section_content:
                            if content_item["type"] == "text":
                                issue["code_sample"] = content_item["value"]
                    
                    elif re.match(r'^NO\.\d+\.\s代码位置$', section_title):
                        code_item = {}
                        
                        if section_content and section_content[0]["type"] == "text":
                            code_location_num = section_content[0]["value"]
                            splitor = code_location_num.split(":")
                            if len(splitor) >= 2:
                                code_item["code_location"] = splitor[0]
                                code_item["code_line_num"] = splitor[1]
                        
                        if "children" in section:
                            for child_section in section["children"]:
                                for content_item in child_section.get("content", []):
                                    if content_item["type"] == "table" and content_item["value"]:
                                        if len(content_item["value"]) > 0 and len(content_item["value"][-1]) > 1:
                                            code_item["code_details"] = content_item["value"][-1][1]
                        
                        if code_item:  # 只添加非空的代码项
                            issue["code_list"].append(code_item)
                
                result[main_title].append(issue)
    
    return result

def sanitize_filename(filename):
    """清理文件名中的特殊字符"""
    return re.sub(r'[<>:"/\\|?*【】（）：/]', '_', filename)

@server.list_tools()
async def handle_list_tools() -> List[types.Tool]:
    """列出可用的SAST报告处理工具"""
    return [
        types.Tool(
            name="convert_sast_docx_to_json",
            description="将SAST报告的docx文档转换为JSON格式，并拆分为单独的漏洞文件",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "SAST报告docx文件的绝对路径或相对路径"
                    }
                },
                "required": ["file_path"]
            }
        ),
        types.Tool(
            name="get_pending_vulnerability_json_files",
            description="获取.scanissuefix目录中所有待处理的漏洞JSON文件(_new.json)",
            inputSchema={
                "type": "object",
                "properties": {
                    "base_dir": {
                        "type": "string",
                        "description": "基础目录路径，默认为当前目录",
                        "default": "."
                    }
                }
            }
        ),
        types.Tool(
            name="generate_csv_report",
            description="从所有已完成的漏洞JSON文件(_finished.json)生成CSV报告",
            inputSchema={
                "type": "object",
                "properties": {
                    "base_dir": {
                        "type": "string", 
                        "description": "基础目录路径，默认为当前目录",
                        "default": "."
                    }
                }
            }
        )
    ]

@server.call_tool()
async def handle_call_tool(
    name: str, arguments: Dict[str, Any]
) -> List[types.TextContent]:
    """
    处理工具执行请求
    """
    try:
        if name == "convert_sast_docx_to_json":
            file_path = arguments["file_path"]
            
            # 支持相对路径和绝对路径
            if not os.path.isabs(file_path):
                file_path = os.path.abspath(file_path)
            
            if not os.path.exists(file_path):
                return [types.TextContent(type="text", text=f"文件不存在: {file_path}")]
            
            logger.info(f"开始处理文件: {file_path}")
            
            # 解析docx为JSON
            all_docx_json = parse_docx_to_json(file_path)
            result = []
            
            for item in all_docx_json:
                if item["title"] in ["四、漏洞详情", "六、代码规范风险详情"]: 
                    transformed = transform_json([item])
                    if item["title"] in transformed:
                        result.extend(transformed[item["title"]])

            if not result:
                return [types.TextContent(type="text", text="未找到有效的漏洞数据（Medium/High级别）")]

            # 确保输出目录存在
            output_dir = ".scanissuefix"
            os.makedirs(output_dir, exist_ok=True)

            saved_files = []
            
            # 处理每个漏洞
            for index, issue in enumerate(result, start=1):
                issue_title = issue["issue_title"]
                code_list = issue["code_list"]
                total_code_entries = len(code_list)

                if total_code_entries > 5:
                    # 分割代码列表
                    chunk_size = 5
                    chunks = [code_list[i:i + chunk_size] for i in range(0, total_code_entries, chunk_size)]
                    
                    for chunk_index, chunk in enumerate(chunks, start=1):
                        issue_copy = issue.copy()
                        issue_copy["code_list"] = chunk

                        # 构建安全的文件名
                        safe_title = sanitize_filename(issue_title)
                        filename = f"{index}_{safe_title}_{chunk_index}_new.json"
                        file_path = os.path.join(output_dir, filename)

                        with open(file_path, 'w', encoding='utf-8') as f:
                            json.dump(issue_copy, f, ensure_ascii=False, indent=2)

                        saved_files.append(filename)
                        logger.info(f"保存文件: {file_path}")
                else:
                    # 单个文件保存
                    safe_title = sanitize_filename(issue_title)
                    filename = f"{index}_{safe_title}_new.json"
                    file_path = os.path.join(output_dir, filename)
                    
                    with open(file_path, 'w', encoding='utf-8') as f:
                        json.dump(issue, f, ensure_ascii=False, indent=2)
                    
                    saved_files.append(filename)
                    logger.info(f"保存文件: {file_path}")
            
            return [types.TextContent(type="text", text=f"成功处理 {len(result)} 个漏洞，保存了 {len(saved_files)} 个文件到 {output_dir} 目录")]

        elif name == "get_pending_vulnerability_json_files":
            base_dir = arguments.get("base_dir", ".")
            output_dir = os.path.join(base_dir, ".scanissuefix")
            
            if not os.path.exists(output_dir):
                return [types.TextContent(type="text", text=f"目录'{output_dir}'不存在")]
            
            all_files = os.listdir(output_dir)
            new_json_files = [os.path.join(output_dir, file) for file in all_files if file.endswith("_new.json")]

            if not new_json_files:
                return [types.TextContent(type="text", text=f"目录'{output_dir}'中没有找到'_new.json'文件")]
            
            return [types.TextContent(type="text", text=f"找到 {len(new_json_files)} 个待处理文件:\n" + "\n".join(new_json_files))]

        elif name == "generate_csv_report":
            base_dir = arguments.get("base_dir", ".")
            output_dir = os.path.join(base_dir, ".scanissuefix")
            
            if not os.path.exists(output_dir):
                return [types.TextContent(type="text", text=f"目录'{output_dir}'不存在")]
            
            all_files = os.listdir(output_dir)
            finished_json_files = [os.path.join(output_dir, file) for file in all_files if file.endswith("_finished.json")]

            if not finished_json_files:
                return [types.TextContent(type="text", text=f"目录'{output_dir}'中没有找到'_finished.json'文件")]
            
            # 准备CSV报告数据
            report_data = []
            processed_files = 0
            
            for file_path in finished_json_files:
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)

                    for code_item in data.get("code_list", []):
                        report_data.append({
                            "漏洞类型": data.get("issue_title", ""),
                            "漏洞等级": data.get("issue_level", ""),
                            "代码位置": code_item.get("code_location", ""),
                            "代码行号": code_item.get("code_line_num", ""),
                            "代码详情": code_item.get("code_details", ""),
                            "修复状态": code_item.get("status", "missed"),
                            "误报概率": code_item.get("false_positive_probability", ""),
                            "误报原因澄清": code_item.get("false_positive_reason", "")
                        })
                    processed_files += 1
                except Exception as e:
                    logger.warning(f"处理文件 {file_path} 时出错: {e}")

            if not report_data:
                return [types.TextContent(type="text", text="没有找到有效的报告数据")]

            # 定义CSV输出文件路径
            csv_output_path = os.path.join(output_dir, "sast_fix_report.csv")

            # 写入CSV文件
            with open(csv_output_path, 'w', newline='', encoding='utf-8') as csv_file:
                fieldnames = [
                    "漏洞类型", 
                    "漏洞等级", 
                    "代码位置", 
                    "代码行号", 
                    "代码详情", 
                    "修复状态", 
                    "误报概率", 
                    "误报原因澄清"
                ]
                writer = csv.DictWriter(csv_file, fieldnames=fieldnames, quoting=csv.QUOTE_MINIMAL)
                writer.writeheader()
                writer.writerows(report_data)

            return [types.TextContent(type="text", text=f"CSV报告生成成功: {csv_output_path}\n处理了 {processed_files} 个文件，生成 {len(report_data)} 条记录")]

        else:
            return [types.TextContent(type="text", text=f"未知工具: {name}")]
            
    except Exception as e:
        error_msg = f"工具 '{name}' 执行失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return [types.TextContent(type="text", text=error_msg)]

async def async_main():
    """Async main function for the MCP server"""
    try:
        logger.info("启动 SAST 修复服务 MCP 服务器...")
        
        async with mcp.server.stdio.stdio_server() as (read_stream, write_stream):
            await server.run(
                read_stream,
                write_stream,
                InitializationOptions(
                    server_name="sast-fixer-mcp",
                    server_version="0.2.0",
                    capabilities=server.get_capabilities(
                        notification_options=NotificationOptions(),
                        experimental_capabilities={},
                    ),
                ),
            )
    except KeyboardInterrupt:
        logger.info("服务器被用户中断")
    except Exception as e:
        logger.error(f"服务器运行时出错: {e}", exc_info=True)
        raise

def main():
    """Main function to run the MCP server"""
    # 运行MCP服务器
    asyncio.run(async_main())

if __name__ == "__main__":
    main()