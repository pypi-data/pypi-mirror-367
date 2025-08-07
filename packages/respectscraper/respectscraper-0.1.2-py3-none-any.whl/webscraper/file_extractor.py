import io
import logging
from typing import Dict

import openpyxl
import PyPDF2
from docx import Document


class FileExtractor:
    """Extracts content from various file types."""

    def __init__(self, config: Dict):
        self.config = config["file_extraction"]
        self.logger = logging.getLogger(__name__)

    def extract_content(self, file_data: bytes, file_extension: str) -> str:
        """
        Extract text content from file data based on file extension.

        Args:
            file_data: Binary file data
            file_extension: File extension (e.g., '.pdf', '.xlsx')

        Returns:
            Extracted text content as string
        """
        try:
            if file_extension.lower() == ".pdf":
                return self._extract_from_pdf(file_data)
            elif file_extension.lower() in [".xlsx", ".xls"]:
                return self._extract_from_excel(file_data)
            elif file_extension.lower() in [".docx", ".doc"]:
                return self._extract_from_word(file_data)
            elif file_extension.lower() == ".txt":
                return self._extract_from_text(file_data)
            else:
                self.logger.warning(f"Unsupported file type: {file_extension}")
                return f"Unsupported file type: {file_extension}"

        except Exception as e:
            self.logger.error(f"Content extraction failed for {file_extension}: {e}")
            return f"Content extraction failed: {str(e)}"

    def _extract_from_pdf(self, file_data: bytes) -> str:
        """Extract text from PDF file."""
        try:
            pdf_file = io.BytesIO(file_data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_content = []
            num_pages = len(pdf_reader.pages)

            self.logger.info(f"Extracting content from PDF with {num_pages} pages")

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    self.logger.warning(
                        f"Failed to extract text from page {page_num + 1}: {e}"
                    )
                    text_content.append(
                        f"--- Page {page_num + 1} ---\n[Text extraction failed: {e}]"
                    )

            extracted_text = "\n\n".join(text_content)

            metadata_info = []
            try:
                if pdf_reader.metadata:
                    if pdf_reader.metadata.title:
                        metadata_info.append(f"Title: {pdf_reader.metadata.title}")
                    if pdf_reader.metadata.author:
                        metadata_info.append(f"Author: {pdf_reader.metadata.author}")
                    if pdf_reader.metadata.subject:
                        metadata_info.append(f"Subject: {pdf_reader.metadata.subject}")
                    if pdf_reader.metadata.creator:
                        metadata_info.append(f"Creator: {pdf_reader.metadata.creator}")
            except Exception as e:
                self.logger.warning(f"Failed to extract PDF metadata: {e}")

            if metadata_info:
                extracted_text = f"""=== PDF METADATA ===\n{chr(10).join(metadata_info)}
                \n=== PDF CONTENT ===\n\n{extracted_text}"""

            return extracted_text

        except Exception as e:
            self.logger.error(f"PDF extraction failed: {e}")
            return f"PDF extraction failed: {str(e)}"

    def _extract_from_excel(self, file_data: bytes) -> str:
        """Extract text from Excel file."""
        try:
            excel_file = io.BytesIO(file_data)
            workbook = openpyxl.load_workbook(excel_file, data_only=True)

            content_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_content = []

                self.logger.info(f"Processing Excel sheet: {sheet_name}")

                if (
                    sheet.max_row == 1
                    and sheet.max_column == 1
                    and sheet.cell(1, 1).value is None
                ):
                    sheet_content.append("[Empty sheet]")
                else:
                    for row in sheet.iter_rows(
                        min_row=1,
                        max_row=sheet.max_row,
                        min_col=1,
                        max_col=sheet.max_column,
                        values_only=True,
                    ):
                        row_values = []
                        for cell_value in row:
                            if cell_value is not None:
                                row_values.append(str(cell_value))
                            else:
                                row_values.append("")

                        if any(val.strip() for val in row_values):
                            sheet_content.append("\t".join(row_values))

                if sheet_content:
                    content_parts.append(
                        f"=== SHEET: {sheet_name} ===\n" + "\n".join(sheet_content)
                    )

            return (
                "\n\n".join(content_parts)
                if content_parts
                else "[No content found in Excel file]"
            )

        except Exception as e:
            self.logger.error(f"Excel extraction failed: {e}")
            return f"Excel extraction failed: {str(e)}"

    def _extract_from_word(self, file_data: bytes) -> str:
        """Extract text from Word document."""
        try:
            word_file = io.BytesIO(file_data)
            document = Document(word_file)

            content_parts = []

            paragraphs = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)

            if paragraphs:
                content_parts.append(
                    "=== DOCUMENT TEXT ===\n" + "\n\n".join(paragraphs)
                )

            tables_content = []
            for table_num, table in enumerate(document.tables):
                table_rows = []
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        row_cells.append(cell.text.strip())
                    if any(cell for cell in row_cells):
                        table_rows.append("\t".join(row_cells))

                if table_rows:
                    tables_content.append(
                        f"--- Table {table_num + 1} ---\n" + "\n".join(table_rows)
                    )

            if tables_content:
                content_parts.append("=== TABLES ===\n" + "\n\n".join(tables_content))

            try:
                core_props = document.core_properties
                metadata_info = []

                if core_props.title:
                    metadata_info.append(f"Title: {core_props.title}")
                if core_props.author:
                    metadata_info.append(f"Author: {core_props.author}")
                if core_props.subject:
                    metadata_info.append(f"Subject: {core_props.subject}")
                if core_props.comments:
                    metadata_info.append(f"Comments: {core_props.comments}")

                if metadata_info:
                    content_parts.insert(
                        0, "=== DOCUMENT METADATA ===\n" + "\n".join(metadata_info)
                    )

            except Exception as e:
                self.logger.warning(f"Failed to extract Word document metadata: {e}")

            return (
                "\n\n".join(content_parts)
                if content_parts
                else "[No content found in Word document]"
            )

        except Exception as e:
            self.logger.error(f"Word document extraction failed: {e}")
            return f"Word document extraction failed: {str(e)}"

    def _extract_from_text(self, file_data: bytes) -> str:
        """Extract text from plain text file."""
        try:
            encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]

            for encoding in encodings:
                try:
                    content = file_data.decode(encoding)
                    self.logger.info(f"Successfully decoded text file using {encoding}")
                    return content
                except UnicodeDecodeError:
                    continue

            content = file_data.decode("utf-8", errors="replace")
            self.logger.warning("Text file decoded with error replacement")
            return content

        except Exception as e:
            self.logger.error(f"Text extraction failed: {e}")
            return f"Text extraction failed: {str(e)}"

    def get_file_info(self, file_data: bytes, file_extension: str) -> Dict:
        """
        Get basic information about the file without extracting full content.

        Args:
            file_data: Binary file data
            file_extension: File extension

        Returns:
            Dict with file information
        """
        info = {
            "size_bytes": len(file_data),
            "size_mb": round(len(file_data) / (1024 * 1024), 2),
            "type": file_extension,
            "extractable": file_extension.lower()
            in self.config["supported_extensions"],
        }

        try:
            if file_extension.lower() == ".pdf":
                pdf_file = io.BytesIO(file_data)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                info["pages"] = len(pdf_reader.pages)
                if pdf_reader.metadata:
                    info["title"] = pdf_reader.metadata.title
                    info["author"] = pdf_reader.metadata.author

            elif file_extension.lower() in [".xlsx", ".xls"]:
                excel_file = io.BytesIO(file_data)
                workbook = openpyxl.load_workbook(excel_file)
                info["sheets"] = list(workbook.sheetnames)
                info["sheet_count"] = len(workbook.sheetnames)

            elif file_extension.lower() in [".docx", ".doc"]:
                word_file = io.BytesIO(file_data)
                document = Document(word_file)
                info["paragraphs"] = len(document.paragraphs)
                info["tables"] = len(document.tables)
                if document.core_properties.title:
                    info["title"] = document.core_properties.title
                if document.core_properties.author:
                    info["author"] = document.core_properties.author

        except Exception as e:
            self.logger.warning(f"Failed to get detailed file info: {e}")
            info["error"] = str(e)

        return info
