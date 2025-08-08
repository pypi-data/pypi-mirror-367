import os
from typing import Any, Optional
from datetime import datetime
import docx
import io

from pilottai_tools.source.base.base_input import BaseInputSource


class DocInput(BaseInputSource):
    """
    Input base for processing Microsoft Word documents (.doc, .docx).
    Extracts and processes text content from Word documents.
    """

    def __init__(
        self,
        name: str,
        file_path: Optional[str] = None,
        file_content: Optional[bytes] = None,
        include_headers: bool = True,
        include_tables: bool = True,
        **kwargs
    ):
        super().__init__(name=name, **kwargs)
        self.file_path = file_path
        self.file_content = file_content
        self.include_headers = include_headers
        self.include_tables = include_tables
        self.text_content = None
        self.doc = None

    async def connect(self) -> bool:
        """Check if the document is accessible"""
        try:
            if self.file_content is not None:
                # Load from binary content
                self.doc = docx.Document(io.BytesIO(self.file_content))
                self.is_connected = True
                return True

            elif self.file_path:
                # Check if file exists
                if not os.path.exists(self.file_path):
                    self.logger.error(f"Document file not found: {self.file_path}")
                    self.is_connected = False
                    return False

                # Check if file is readable
                if not os.access(self.file_path, os.R_OK):
                    self.logger.error(f"Document file not readable: {self.file_path}")
                    self.is_connected = False
                    return False

                # Try to load the document
                self.doc = docx.Document(self.file_path)
                self.is_connected = True
                return True

            self.logger.error("No file path or content provided")
            self.is_connected = False
            return False

        except Exception as e:
            self.logger.error(f"Connection error: {str(e)}")
            self.is_connected = False
            return False

    async def query(self, query: str) -> Any:
        """Search for query in the document content"""
        if not self.is_connected or not self.text_content:
            if not await self.extract_text():
                raise ValueError("Could not extract text from document")

        self.access_count += 1
        self.last_access = datetime.now()

        # Simple search implementation
        results = []
        if query.lower() in self.text_content.lower():
            context_size = 200  # Characters before and after match

            # Find all occurrences
            start_idx = 0
            query_lower = query.lower()
            text_lower = self.text_content.lower()

            while True:
                idx = text_lower.find(query_lower, start_idx)
                if idx == -1:
                    break

                # Get context around the match
                context_start = max(0, idx - context_size)
                context_end = min(len(self.text_content), idx + len(query) + context_size)
                context = self.text_content[context_start:context_end]

                results.append({
                    "match": self.text_content[idx:idx + len(query)],
                    "context": context,
                    "position": idx
                })

                start_idx = idx + len(query)

        return results

    async def validate_content(self) -> bool:
        """Validate that document content can be extracted"""
        if not self.is_connected:
            if not await self.connect():
                return False

        return await self.extract_text()

    async def extract_text(self) -> bool:
        """Extract text content from the document"""
        try:
            if not self.doc:
                if not await self.connect():
                    return False

            paragraphs = []

            # Extract text from paragraphs
            for para in self.doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract text from tables if requested
            if self.include_tables:
                for table in self.doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(cell.text for cell in row.cells)
                        if row_text.strip():
                            paragraphs.append(row_text)

            # Extract text from headers and footers if requested
            if self.include_headers:
                for section in self.doc.sections:
                    # Headers
                    for header in [section.header]:
                        for para in header.paragraphs:
                            if para.text.strip():
                                paragraphs.append(f"[Header] {para.text}")

                    # Footers
                    for footer in [section.footer]:
                        for para in footer.paragraphs:
                            if para.text.strip():
                                paragraphs.append(f"[Footer] {para.text}")

            self.text_content = '\n'.join(paragraphs)
            return bool(self.text_content)

        except Exception as e:
            self.logger.error(f"Error extracting text from document: {str(e)}")
            return False

    async def _process_content(self) -> None:
        """Process and chunk the document content"""
        if not self.text_content:
            if not await self.extract_text():
                return

        self.chunks = self._chunk_text(self.text_content)
        source_desc = self.file_path if self.file_path else "binary content"
        self.logger.info(f"Created {len(self.chunks)} chunks from document {source_desc}")
